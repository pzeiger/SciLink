"""
Tool definitions and schemas for the AnalysisOrchestratorAgent.
Supports both OpenAI (JSON schemas) and LiteLLM formats.

Each analysis run creates a unique output directory to ensure traceability
and prevent output collisions when analyzing multiple datasets.

Per-file JSON sidecar metadata
------------------------------
When a data directory contains JSON files whose stems match data files
(e.g. ``spec_5K.csv`` ↔ ``spec_5K.json``), they are treated as *sidecar
metadata* rather than global metadata.  ``run_analysis`` will attempt to
extract the series control variable from the sidecars automatically via
LLM reasoning.  If extraction fails, the user is prompted and shown the
sidecar contents to help them specify the variable manually.
"""

import hashlib
import json
import logging
import re
import time
import numpy as np
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Callable, List

from .metadata_converter import (
    generate_metadata_json_from_text,
    METADATA_SCHEMA_DICT,
    check_schema_conformance,
    normalize_metadata_dict,
    normalize_metadata_dict_with_llm,
)
from ..lit_agents import OwlLiteratureAgent, NoveltyScorer, FittingModelLiteratureAgent
from ..lit_agents.optimize_query_for_analysis import optimize_query_for_analysis
from .recommendation_agent import RecommendationAgent
from .feature_table import write_feature_table
from ...skills.loader import list_skills, list_all_skills, load_skill
# Note: DFTOrchestrator is imported lazily inside `run_dft_workflow` to avoid
# pulling in the optional [sim] extras (ase, atomate2, pymatgen) on every
# AnalysisOrchestratorAgent instantiation.


# Full-text extraction for the read_document tool — a few documents read
# straight into the LLM context, with no embeddings / chunking / vector store
# (that is the planning KB's job, for large corpora). The fitz / docx readers
# are imported lazily so this module stays importable without them.
_READ_DOC_MAX_CHARS = 200_000  # ~50k tokens; longer documents are truncated


def _extract_document_text(path: Path) -> Dict[str, Any]:
    """Extract plain text from a PDF / DOCX / Markdown / text file.

    Returns a dict with ``text`` plus metadata (page/paragraph count,
    ``n_chars``, ``truncated``). Raises ValueError for an unsupported
    extension; reader errors propagate to the caller.
    """
    ext = path.suffix.lower()
    info: Dict[str, Any] = {}
    if ext == ".pdf":
        import fitz
        doc = fitz.open(path)
        try:
            info["n_pages"] = doc.page_count
            text = "\n\n".join(
                doc[i].get_text() or "" for i in range(doc.page_count)
            )
        finally:
            doc.close()
    elif ext == ".docx":
        import docx
        d = docx.Document(str(path))
        info["n_paragraphs"] = len(d.paragraphs)
        text = "\n".join(p.text for p in d.paragraphs)
    elif ext in (".md", ".txt"):
        text = path.read_text(errors="replace")
    else:
        raise ValueError(
            f"Unsupported document type '{ext}' — read_document handles "
            f".pdf, .docx, .md, and .txt."
        )
    text = text.strip()
    info["truncated"] = len(text) > _READ_DOC_MAX_CHARS
    if info["truncated"]:
        text = text[:_READ_DOC_MAX_CHARS]
    info["text"] = text
    info["n_chars"] = len(text)
    return info


def _build_skill_description(agent_registry: dict = None,
                              custom_skills: dict = None) -> str:
    """Build the ``skill`` parameter description for ``run_analysis``.

    Auto-discovers built-in skill domains and inspects the agent registry
    to determine which agents accept a ``skill`` parameter.
    """
    import inspect

    parts = [
        "Domain skill name or path to a custom .md skill file. May be a "
        "single string or a list of strings to load multiple skills at once "
        "(useful for cross-domain tasks). For ImageAnalysisAgent, omit this "
        "unless the user explicitly requests a specific skill — the agent "
        "inspects the actual image and auto-selects a skill if one is relevant."
    ]

    # Discover which agents support skills from their analyze() signature.
    # Registry entries use lazy class loading — resolve class_path if needed.
    if agent_registry:
        supported = []
        for entry in agent_registry.values():
            cls = entry.get("class")
            if cls is None and "class_path" in entry:
                try:
                    module_path, cls_name = entry["class_path"].rsplit(".", 1)
                    import importlib
                    mod = importlib.import_module(module_path)
                    cls = getattr(mod, cls_name)
                except Exception:
                    continue
            if cls is None:
                continue
            try:
                sig = inspect.signature(cls.analyze)
                if "skill" in sig.parameters:
                    supported.append(entry["name"])
            except (ValueError, TypeError):
                continue
        if supported:
            parts.append(f"Supported by: {', '.join(supported)}.")

    # Auto-discover all built-in skill domains with descriptions.
    # Prefer the frontmatter `description` field when present; fall back
    # to the first line of the overview section.
    for domain, names in list_all_skills().items():
        skill_descs = []
        for name in names:
            try:
                parsed = load_skill(name, domain=domain)
                desc = (parsed.get("meta") or {}).get("description")
                if not desc:
                    desc = parsed.get("overview", "").split("\n")[0].strip()
                # Trim trailing punctuation so the join with ". " below
                # doesn't produce ".." or ".;".
                desc = desc.rstrip(".;,") if desc else desc
                skill_descs.append(f"'{name}' — {desc}" if desc else f"'{name}'")
            except Exception:
                skill_descs.append(f"'{name}'")
        parts.append(f"Built-in {domain} skills: {'; '.join(skill_descs)}.")

    if custom_skills:
        parts.append(f"Custom skills: {sorted(custom_skills.keys())}.")

    return " ".join(parts)


# Names that are always treated as global (directory-level) metadata files,
# never as per-file sidecars, even if their stem happens to match a data file.
_GLOBAL_METADATA_NAMES = frozenset([
    "metadata.json", "meta.json", "info.json", "experiment.json",
])


def _detect_sidecar_jsons(
    data_files: list[Path],
    all_files: list[Path],
) -> tuple[dict[str, Path], list[Path]]:
    """Identify JSON files that are stem-matched sidecars for data files.

    A JSON file is a *sidecar* when its stem matches a data file's stem
    (e.g. ``spec_5K.json`` ↔ ``spec_5K.csv``).  Files whose names are in
    ``_GLOBAL_METADATA_NAMES`` are always treated as global metadata.

    Returns
    -------
    sidecar_map : dict[str, Path]
        ``{data_filename: sidecar_Path}`` for every matched pair.
    global_jsons : list[Path]
        JSON files that are **not** sidecars (global metadata or unmatched).
    """
    data_stems = {f.stem: f for f in data_files}
    json_files = [f for f in all_files if f.suffix.lower() == ".json"]

    sidecar_map: dict[str, Path] = {}
    global_jsons: list[Path] = []

    for jf in json_files:
        if jf.name.lower() in _GLOBAL_METADATA_NAMES:
            global_jsons.append(jf)
        elif jf.stem in data_stems:
            sidecar_map[data_stems[jf.stem].name] = jf
        else:
            global_jsons.append(jf)

    return sidecar_map, global_jsons


def _parse_key_list_response(raw: str, valid_keys: list[str]) -> list[str]:
    """Extract an ordered list of key names from a (possibly noisy) LLM
    response. The LLM is asked for the control-variable field names
    comma-separated, most primary first; ``NONE`` means none qualifies.

    Returns the matched keys in response order (deduped), or ``[]``.
    """
    # Strip markdown fences
    text = raw.strip()
    if text.startswith("```"):
        text = text.split("```")[1].split("```")[0].strip()

    if text.strip().strip("\"'`*. ").upper() == "NONE":
        return []

    # Ideal case: comma / newline / semicolon separated tokens.
    ordered: list[str] = []
    for token in re.split(r"[,\n;]+", text):
        tok = token.strip().strip("\"'`*. ")
        if tok in valid_keys and tok not in ordered:
            ordered.append(tok)
    if ordered:
        return ordered

    # Fallback: scan for any valid key as a whole word, in order of appearance.
    positions: list[tuple[int, str]] = []
    for key in valid_keys:
        m = re.search(rf"\b{re.escape(key)}\b", text)
        if m:
            positions.append((m.start(), key))
    positions.sort()
    return [k for _, k in positions]


def _llm_identify_control_variables(
    varying_keys: list[str],
    sidecar_data: dict[str, dict],
    model,
    logger: logging.Logger,
    experimental_context: dict | None = None,
) -> list[str]:
    """Use the LLM to identify which varying sidecar fields are genuine
    independent control variables, ordered most-primary-first.

    An experiment may deliberately vary several variables at once (a
    factorial / grid design). The LLM returns the full set, so the caller
    can record a primary axis plus secondary variables rather than being
    forced into a single pick.

    Parameters
    ----------
    varying_keys : list[str]
        Numeric sidecar keys whose values differ across files (≥ 1).
    sidecar_data : dict[str, dict]
        ``{data_filename: sidecar_dict}`` for every file.
    model
        LLM wrapper with a ``generate_content`` method.
    logger
        Logger instance.
    experimental_context : dict | None
        Optional dict with keys ``"objective"``, ``"hints"``, and/or
        ``"metadata"`` providing broader experimental context.

    Returns
    -------
    list[str]
        Control-variable key names, primary first; ``[]`` when none of the
        candidates is a genuine control variable.
    """
    # Build a summary of each candidate key and its per-file values
    candidates_summary = {}
    for key in varying_keys:
        candidates_summary[key] = {
            fname: d[key] for fname, d in sidecar_data.items()
        }

    # Assemble context lines
    context_parts = []
    if experimental_context:
        if experimental_context.get("objective"):
            context_parts.append(
                f"Analysis objective: {experimental_context['objective']}"
            )
        if experimental_context.get("hints"):
            context_parts.append(
                f"User hints: {experimental_context['hints']}"
            )
        meta = experimental_context.get("metadata")
        if isinstance(meta, dict):
            # Include high-level experiment info, not the full blob
            for k in ("experiment_type", "experiment", "sample"):
                if k in meta:
                    context_parts.append(f"{k}: {json.dumps(meta[k])}")

    context_block = "\n".join(context_parts) if context_parts else "None provided."

    keys_list = ", ".join(varying_keys)
    prompt = (
        "You are a scientific data analysis assistant. A user has a series of "
        "data files, each accompanied by a JSON sidecar containing per-file "
        "metadata. The following numeric fields change across the files.\n\n"
        "Identify which of these fields are genuine **independent control "
        "variables** — physical or experimental quantities the experimenter "
        "intentionally varied across measurements (e.g. temperature, "
        "concentration, voltage, pressure, dose, time).\n\n"
        "An experiment may deliberately vary MORE THAN ONE at once — a "
        "factorial / grid design varies several parameters across the "
        "measurements. Report ALL genuine control variables; do NOT force a "
        "single choice and do NOT answer 'uncertain' just because there are "
        "several.\n\n"
        "For in-situ, time-resolved, or kinetic experiments, elapsed time "
        "or total time IS a control variable — do not dismiss it as mere "
        "acquisition metadata. "
        "Instrument and acquisition parameters that happen to differ "
        "(e.g. laser_power, integration_time, slit_width, probe_current) "
        "are NOT control variables — exclude them. Note: 'integration_time' "
        "(detector exposure per scan) is an acquisition setting, but 'total "
        "time' or 'elapsed time' (cumulative experiment duration) is a real "
        "control variable.\n\n"
        "It is also possible that NONE of the listed fields is a true "
        "control variable — for example the real one was set manually and "
        "not recorded in the sidecar metadata.\n\n"
        f"Experimental context:\n{context_block}\n\n"
        f"Candidate fields and their per-file values:\n"
        f"{json.dumps(candidates_summary, indent=2)}\n\n"
        "RESPONSE FORMAT: a single line, comma-separated, listing the control "
        "variable field names MOST PRIMARY FIRST (the dominant experimental "
        "axis first, the rest after). No explanations, no other text. If "
        "none of the fields is a genuine control variable, respond with "
        "exactly NONE.\n"
        f"Choose only from: {keys_list}\n"
    )

    try:
        response = model.generate_content(contents=[prompt])
        raw = (
            response.text
            if hasattr(response, "text")
            else str(response)
        ).strip()

        chosen = _parse_key_list_response(raw, varying_keys)

        if not chosen:
            logger.info(
                "LLM identified no series control variable "
                "(response: %r, candidates: %s)",
                raw,
                varying_keys,
            )
            return []

        logger.info(
            "LLM identified control variable(s) %s from candidates %s",
            chosen,
            varying_keys,
        )
        return chosen

    except Exception as exc:
        logger.warning("LLM control-variable identification failed: %s", exc)
        return []


# Keys produced by the LLM metadata normalization schema.
_CANONICAL_SCHEMA_KEYS = frozenset({
    "experiment_type", "experiment", "sample", "spatial_info",
    "energy_range", "title", "data_columns", "xlabel", "ylabel",
    "custom_processing_instruction",
})

# Keys managed by the sidecar / series extraction pipeline.
_INTERNAL_KEYS = frozenset({"per_file_metadata", "series"})


def _structure_metadata_for_save(metadata: dict) -> dict:
    """Restructure flat current_metadata into grouped sections for saving.

    The runtime ``current_metadata`` dict is kept flat so that agents can
    access keys directly (``system_info.get("title")``).  For the saved
    ``metadata_used.json`` we reorganise into three clear groups:

    * **global** — normalised experiment-level fields from the canonical
      schema (experiment, sample, title, xlabel, …).
    * **per_file_metadata** / **series** — kept at top level as-is.
    * **raw_instrument** — remaining passthrough fields from sidecar
      synthesis that were not consumed by the LLM normalisation.
    """
    if not isinstance(metadata, dict):
        return metadata

    global_section: dict = {}
    raw_section: dict = {}
    internal_section: dict = {}

    for key, value in metadata.items():
        if key in _INTERNAL_KEYS:
            internal_section[key] = value
        elif key in _CANONICAL_SCHEMA_KEYS:
            global_section[key] = value
        else:
            raw_section[key] = value

    # Build result with global first for readability
    result: dict = {}
    if global_section:
        result["global"] = global_section
    result.update(internal_section)
    if raw_section:
        result["raw_instrument"] = raw_section

    return result


def _extract_series_from_sidecars(
    sidecar_map: dict[str, Path],
    data_files: list[Path],
    logger: logging.Logger,
    model=None,
    experimental_context: dict | None = None,
) -> tuple[dict | None, dict[str, dict]]:
    """Try to auto-build series metadata from per-file sidecar JSONs.

    The algorithm:

    1. Load every sidecar; bail out if coverage is incomplete.
    2. Collect top-level numeric keys common to **all** sidecars.
    3. Keep only keys whose values **differ** across files — this naturally
       eliminates constant instrument settings (e.g. integration_time=1.0
       in every file).
    4. If *model* is provided, ask the LLM to evaluate the remaining
       candidates (even if there is only one) and decide whether any is a
       true independent control variable.  A single varying key might still
       be just an acquisition setting; the real control variable may not be
       recorded in the sidecars at all.
    5. Fall back to ``None`` (user prompt) if no model or the LLM cannot
       decide.

    Returns
    -------
    series_meta : dict | None
        ``{"variable": ..., "values": {fname: val}, "unit": ...}`` for the
        primary control variable, plus an optional ``"secondary_variables"``
        list of the same shape for any others that co-vary (grid designs).
        ``None`` when extraction is not possible.
    per_file_meta : dict[str, dict]
        Full sidecar contents keyed by data filename (always returned,
        even when ``series_meta`` is ``None``).
    """
    per_file_meta: dict[str, dict] = {}
    sidecar_data: dict[str, dict] = {}

    # 1. Load all sidecars
    for fname, jpath in sidecar_map.items():
        try:
            with open(jpath, "r") as f:
                content = json.load(f)
            if isinstance(content, dict):
                sidecar_data[fname] = content
                per_file_meta[fname] = content
        except (json.JSONDecodeError, OSError) as exc:
            logger.warning("Failed to load sidecar %s: %s", jpath.name, exc)

    if not sidecar_data:
        return None, per_file_meta

    # 2. Only proceed when every data file has a sidecar
    data_fnames = {f.name for f in data_files}
    if set(sidecar_data.keys()) != data_fnames:
        logger.info(
            "Sidecar coverage incomplete: %d/%d files",
            len(sidecar_data),
            len(data_fnames),
        )
        return None, per_file_meta

    # 3. Common top-level numeric keys across all sidecars
    all_keys: set[str] | None = None
    for d in sidecar_data.values():
        numeric_keys = {
            k for k, v in d.items() if isinstance(v, (int, float))
        }
        all_keys = numeric_keys if all_keys is None else all_keys & numeric_keys

    if not all_keys:
        return None, per_file_meta

    # 4. Keep only keys whose values differ across files
    varying_keys = []
    for key in all_keys:
        vals = [d[key] for d in sidecar_data.values()]
        if len(set(vals)) > 1:
            varying_keys.append(key)

    if not varying_keys:
        return None, per_file_meta

    # 5. Ask the LLM which varying fields are genuine control variables
    #    (even a single one could be just an acquisition setting; a grid
    #    design may have several).
    if model is not None:
        print(
            f"    Varying fields in sidecars: {varying_keys}. "
            f"Asking LLM to identify the control variable(s)..."
        )
        control_vars = _llm_identify_control_variables(
            varying_keys, sidecar_data, model, logger,
            experimental_context=experimental_context,
        )
    else:
        control_vars = []

    if not control_vars:
        logger.info(
            "Could not identify a series control variable from "
            "sidecar candidates: %s",
            varying_keys,
        )
        return None, per_file_meta

    # 6. Build series metadata. The primary control variable becomes
    #    variable/values/unit (the axis used for file ordering, scouting and
    #    trend analysis); any others are recorded as secondary_variables so
    #    the analysis is told the full set of conditions that co-vary.
    def _meta_for(var: str) -> dict:
        vals = {fname: sidecar_data[fname][var] for fname in sidecar_data}
        unit = ""
        sample_sidecar = next(iter(sidecar_data.values()))
        for unit_key in (f"{var}_unit", f"{var}_units", "unit", "units"):
            if unit_key in sample_sidecar:
                unit = str(sample_sidecar[unit_key])
                break
        return {"variable": var, "values": vals, "unit": unit}

    series_meta = _meta_for(control_vars[0])
    if len(control_vars) > 1:
        series_meta["secondary_variables"] = [
            _meta_for(v) for v in control_vars[1:]
        ]
        logger.info(
            "Series control variables: primary=%s, secondary=%s",
            control_vars[0], control_vars[1:],
        )

    return series_meta, per_file_meta


class AnalysisOrchestratorTools:
    """
    Manages tool definitions, schemas, and execution for the AnalysisOrchestratorAgent.
    """

    def __init__(self, orchestrator_instance):
        """
        Args:
            orchestrator_instance: Reference to the parent AnalysisOrchestratorAgent
        """
        self.orch = orchestrator_instance
        self.logger = logging.getLogger(self.__class__.__name__)

        # Agent display dicts — populated from the orchestrator's live registry
        # so that custom agents registered via register_agent() appear here too.
        self.AGENT_NAMES: Dict[int, str] = {}
        self.AGENT_DESCRIPTIONS: Dict[int, str] = {}
        self._sync_from_registry()

        # Build function map and schemas
        self.functions_map: Dict[str, Callable] = {}
        self.openai_schemas: list = []

        self._register_all_tools()

    def _sync_from_registry(self) -> None:
        """Rebuild AGENT_NAMES and AGENT_DESCRIPTIONS from the orchestrator's registry."""
        registry = getattr(self.orch, "_agent_registry", {})
        self.AGENT_NAMES = {aid: e["name"] for aid, e in registry.items()}
        self.AGENT_DESCRIPTIONS = {aid: e["description"] for aid, e in registry.items()}

    def _get_human_feedback_enabled(self) -> bool:
        """Get current human feedback setting from orchestrator."""
        return getattr(self.orch, '_enable_human_feedback', True)

    def _check_instruction_redundancy(self, existing: str, new: str) -> bool:
        """Ask the LLM whether two preprocessing instructions are redundant.

        Returns ``True`` if the instructions describe essentially the same
        processing operation (even if worded differently), ``False`` if they
        are genuinely distinct steps.  Falls back to ``False`` on any error
        so that the caller never blocks on a failed check.
        """
        prompt = (
            "You are a scientific data processing expert. Determine whether "
            "the following two preprocessing instructions describe the SAME "
            "operation (even if worded differently) or genuinely DIFFERENT "
            "processing steps.\n\n"
            f"Instruction A: {existing}\n\n"
            f"Instruction B: {new}\n\n"
            "If they are essentially the same operation (e.g. both describe "
            "baseline division, both describe the same normalization, etc.), "
            "respond with exactly: REDUNDANT\n"
            "If they are genuinely different processing steps, respond with "
            "exactly: DISTINCT\n\n"
            "Respond with a SINGLE word — REDUNDANT or DISTINCT."
        )
        try:
            response = self.orch.model.generate_content(contents=[prompt])
            raw = (
                response.text if hasattr(response, "text") else str(response)
            ).strip().upper()
            return "REDUNDANT" in raw
        except Exception as e:
            self.logger.warning(f"LLM redundancy check failed: {e}")
            return False

    def _replace_metadata(self, new_metadata: dict) -> dict | None:
        """Replace current_metadata, preserving any custom_processing_instruction.

        If the old metadata had a custom_processing_instruction that would be
        lost (i.e. the new metadata doesn't include one), it is carried
        forward automatically.

        Returns a warning dict if both old and new metadata contain
        *different* custom_processing_instructions (the new one wins, but
        the caller should surface the conflict).  Returns ``None`` otherwise.
        """
        _CPI = "custom_processing_instruction"
        old_instruction = None
        if self.orch.current_metadata:
            old_instruction = self.orch.current_metadata.get(_CPI)

        self.orch.current_metadata = new_metadata

        if not old_instruction:
            return None  # Nothing to preserve

        new_instruction = new_metadata.get(_CPI)

        if not new_instruction:
            # New metadata has no instruction — carry forward the old one
            new_metadata[_CPI] = old_instruction
            return None

        # Both have instructions — keep the new one but warn
        if old_instruction.strip() != new_instruction.strip():
            return {
                "preprocessing_warning": (
                    "New metadata contains a custom_processing_instruction that "
                    "differs from a previously set one. The new instruction is "
                    "being used. Review to ensure correctness."
                ),
                "previous_instruction": old_instruction,
                "current_instruction": new_instruction,
            }
        return None

    def _examine_hdf5(self, path: Path, result: dict) -> None:
        """Populate *result* with shape/dimensions/metadata for an
        HDF5 file.

        Uses the official ``nexusformat`` library for NeXus-conformant
        files (handles NXentry/NXdata/signal/axes per the standard,
        including older conventions).  Falls back to a generic ``h5py``
        walk for non-NeXus HDF5 so the LLM still gets useful structural
        info (group/dataset shapes, dtypes, attrs).
        """
        try:
            import h5py  # noqa: F401  (used in the fallback walk)
        except ImportError:
            result["data_type"] = "unknown"
            result["suggested_agents"] = []
            result["message"] = (
                "Reading HDF5 files requires h5py. Install with: "
                "pip install h5py"
            )
            return

        # Try NeXus parse first.
        parsed = None
        nexus_error = None
        try:
            import nexusformat.nexus as nx
            try:
                nx_root = nx.nxload(str(path), mode="r")
                parsed = self._parse_nexus(nx_root)
            except Exception as exc:
                nexus_error = str(exc)
        except ImportError:
            nexus_error = (
                "nexusformat not installed; falling back to a generic "
                "HDF5 walk. Install with: pip install nexusformat"
            )

        # Always run the generic content harvester — it surfaces auxiliary
        # metadata groups (sidpy_metadata, hyperspy metadata, …) that the
        # NeXus standards-aware parse doesn't touch, and is the bridge
        # source for ``convert_metadata`` when an h5 happens to embed
        # producer-specific metadata.  Producer-agnostic; bounded payload.
        try:
            harvested = self._harvest_dataset_contents(path)
        except Exception as exc:
            harvested = None
            result.setdefault("hdf5_walk_error", str(exc))

        if parsed is not None:
            self._apply_nexus_to_result(parsed, result)
            if harvested is not None:
                # Don't duplicate the signal/axis arrays — the NeXus parse
                # already surfaced shape/dtype/units for the primary
                # data.  Auxiliary content (everything else) is what's
                # genuinely new here.
                result["root_attrs"] = harvested["root_attrs"]
                result["hdf5_datasets"] = harvested["datasets"][:50]
                result["dataset_count"] = harvested["dataset_count"]
                if harvested["truncated"]:
                    result["hdf5_content_truncated"] = True
            return

        # Generic fallback — no NeXus structure.
        if nexus_error:
            result["nexus_parse_error"] = nexus_error
        if harvested is not None:
            self._apply_harvested_to_result(harvested, result)
        else:
            result["data_type"] = "unknown"
            result["suggested_agents"] = []
            result.setdefault("message", "Failed to read HDF5")

    @staticmethod
    def _parse_nexus(nx_root) -> "dict | None":
        """Extract NeXus-standard structure (signal, axes, dimensions)
        from an HDF5 file via ``nexusformat``'s standards-aware accessors.

        Producer-specific metadata layouts (sidpy's ``sidpy_metadata``
        group, HyperSpy's ``signal.metadata`` tree, …) are *not*
        interpreted here — bridging third-party metadata dialects to
        SciLink's canonical schema is the job of ``convert_metadata``,
        which is producer-agnostic.

        Returns ``None`` if the file has no NXentry/NXdata or no signal.
        """
        entries = list(nx_root.NXentry)
        if not entries:
            return None
        default_entry = nx_root.attrs.get("default")
        entry = next(
            (e for e in entries if e.nxname == default_entry),
            entries[0],
        )

        nxdatas = list(entry.NXdata)
        if not nxdatas:
            return None
        default_nxdata = entry.attrs.get("default")
        nxdata = next(
            (d for d in nxdatas if d.nxname == default_nxdata),
            nxdatas[0],
        )

        signal = nxdata.nxsignal
        if signal is None:
            return None

        def _attr(field, key, default=""):
            v = field.attrs.get(key, default)
            if isinstance(v, bytes):
                return v.decode("utf-8", errors="replace")
            return v

        title = _attr(signal, "title") or _attr(signal, "long_name") or ""
        units = _attr(signal, "units") or ""

        try:
            axes = list(nxdata.nxaxes or [])
        except Exception:
            axes = []

        dims_info = []
        for i, dim_len in enumerate(signal.shape):
            entry_d = {"index": i, "length": int(dim_len)}
            if i < len(axes):
                ax = axes[i]
                ax_name = getattr(ax, "nxname", None)
                if ax_name:
                    entry_d["name"] = ax_name
                entry_d["units"] = _attr(ax, "units") or ""
                try:
                    vals = np.asarray(ax.nxdata)
                    if vals.size >= 2:
                        diffs = np.diff(vals)
                        if diffs.size and np.allclose(diffs, diffs[0]):
                            entry_d["start"] = float(vals[0])
                            entry_d["stop"] = float(vals[-1])
                            entry_d["step"] = float(diffs[0])
                except Exception:
                    pass
            dims_info.append(entry_d)

        return {
            "shape": list(signal.shape),
            "dtype": str(signal.dtype),
            "title": title,
            "units": units,
            "dimensions": dims_info,
        }

    @staticmethod
    def _apply_nexus_to_result(parsed: dict, result: dict) -> None:
        """Copy the parsed NeXus payload into the examine_data result and
        derive a shape-based agent suggestion."""
        shape = parsed["shape"]

        result["shape"] = shape
        result["dtype"] = parsed["dtype"]
        result["title"] = parsed["title"]
        result["units"] = parsed["units"]
        result["dimensions"] = parsed["dimensions"]

        if len(shape) == 1:
            result["data_type"] = "1d_data"
            result["suggested_agents"] = [0]
            result["primary_suggestion"] = 0
        elif len(shape) == 2:
            result["data_type"] = "image"
            result["suggested_agents"] = [1]
            result["primary_suggestion"] = 1
        elif len(shape) == 3:
            result["data_type"] = "hyperspectral"
            result["suggested_agents"] = [2]
            result["primary_suggestion"] = 2
        else:
            result["data_type"] = "nd_data"
            result["suggested_agents"] = []

        result["note"] = (
            f"NeXus dataset: shape={shape}, "
            f"{len(parsed['dimensions'])} dimension(s)"
        )

    # Per-dataset and total caps for embedded text/JSON content surfaced
    # to the LLM.  Numeric arrays are *never* included — only their shape
    # and dtype.  Object-dtype contents that don't decode cleanly as
    # UTF-8 are dropped (avoids pickle hazard and binary mojibake).
    _HDF5_PER_DATASET_CONTENT_CAP = 32 * 1024
    _HDF5_TOTAL_CONTENT_BUDGET = 128 * 1024

    @staticmethod
    def _decode_h5_value(v):
        """Best-effort decode of an h5py attribute or scalar value into a
        JSON-serialisable Python object."""
        if isinstance(v, bytes):
            try:
                return v.decode("utf-8")
            except UnicodeDecodeError:
                return v.decode("utf-8", errors="replace")
        if isinstance(v, np.ndarray):
            if v.dtype.kind in ("S", "O"):
                return [
                    item.decode("utf-8", errors="replace")
                    if isinstance(item, bytes) else item
                    for item in v.tolist()
                ]
            return v.tolist()
        if isinstance(v, np.generic):
            return v.item()
        return v

    @classmethod
    def _try_decode_dataset_content(cls, raw):
        """Decode a dataset value to a JSON object or text string, or
        return ``None`` if it can't be safely decoded.

        Refuses bytes that don't decode as UTF-8 (avoids surfacing
        pickled Python objects or binary blobs to the LLM).
        """
        # Bytes scalar: try UTF-8 strictly (no replace — we want to know
        # if it's actually text).
        if isinstance(raw, bytes):
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                return None
        elif isinstance(raw, str):
            text = raw
        elif isinstance(raw, np.ndarray):
            if raw.dtype.kind in ("S", "O"):
                pieces = []
                items = raw.tolist() if raw.shape else [raw.item()]
                if not isinstance(items, list):
                    items = [items]
                for item in items:
                    if isinstance(item, bytes):
                        try:
                            pieces.append(item.decode("utf-8"))
                        except UnicodeDecodeError:
                            return None
                    elif isinstance(item, str):
                        pieces.append(item)
                    else:
                        return None
                text = "\n".join(pieces)
            elif raw.dtype.kind == "U":
                items = raw.tolist() if raw.shape else [raw.item()]
                text = "\n".join(items if isinstance(items, list) else [items])
            else:
                return None  # numeric — caller skips
        else:
            return None

        # Try JSON parse first; if it parses, the structured form is more
        # useful to the LLM than a giant text blob.
        try:
            return json.loads(text)
        except (json.JSONDecodeError, ValueError):
            return text

    @classmethod
    def _harvest_dataset_contents(cls, path: Path) -> dict:
        """Walk an HDF5 file and return its structure plus the contents
        of small text/JSON datasets, bounded by per-dataset and total
        byte budgets.

        Returns
        -------
        dict with keys: ``root_attrs``, ``datasets`` (list of descriptors,
        each with ``path/shape/dtype/attrs`` and optionally ``content``),
        ``dataset_count`` (total found), ``truncated`` (True if budget hit).
        """
        import h5py

        per_cap = cls._HDF5_PER_DATASET_CONTENT_CAP
        total_cap = cls._HDF5_TOTAL_CONTENT_BUDGET
        spent = 0
        truncated = False
        descriptors: list[dict] = []

        with h5py.File(str(path), "r") as f:
            root_attrs = {
                k: cls._decode_h5_value(v) for k, v in f.attrs.items()
            }

            def visit(name, obj):
                nonlocal spent, truncated
                if not isinstance(obj, h5py.Dataset):
                    return
                d = {
                    "path": name,
                    "shape": list(obj.shape),
                    "dtype": str(obj.dtype),
                    "attrs": {
                        k: cls._decode_h5_value(v) for k, v in obj.attrs.items()
                    },
                }

                # Only consider surfacing contents for text/object dtypes.
                # Numeric arrays are intentionally skipped (would dump
                # signal data).
                if obj.dtype.kind in ("S", "U", "O"):
                    # Estimate byte size; refuse large datasets entirely
                    # before reading them.
                    try:
                        nbytes = int(obj.nbytes) if obj.size else 0
                    except Exception:
                        nbytes = 0
                    if nbytes <= per_cap and spent < total_cap:
                        try:
                            raw = obj[()]
                            decoded = cls._try_decode_dataset_content(raw)
                        except Exception:
                            decoded = None

                        if decoded is not None:
                            # Approximate cost as JSON-encoded length so
                            # the budget tracks LLM token consumption.
                            try:
                                cost = len(
                                    json.dumps(decoded, default=str).encode("utf-8")
                                )
                            except Exception:
                                cost = per_cap  # be conservative
                            if cost <= per_cap and spent + cost <= total_cap:
                                d["content"] = decoded
                                spent += cost
                            else:
                                truncated = True

                descriptors.append(d)

            f.visititems(visit)

        return {
            "root_attrs": root_attrs,
            "datasets": descriptors,
            "dataset_count": len(descriptors),
            "truncated": truncated,
        }

    @staticmethod
    def _apply_harvested_to_result(harvested: dict, result: dict) -> None:
        """Populate the examine_data result from a generic harvest (used
        when no NeXus structure was found)."""
        datasets = harvested["datasets"]
        result["root_attrs"] = harvested["root_attrs"]
        result["hdf5_datasets"] = datasets[:50]
        result["dataset_count"] = harvested["dataset_count"]
        if harvested["truncated"]:
            result["hdf5_content_truncated"] = True
        result["note"] = (
            f"HDF5 file with {harvested['dataset_count']} dataset(s); "
            "non-NeXus or unparseable as NeXus."
        )

        # Best-effort agent suggestion from a single top-level dataset.
        if harvested["dataset_count"] == 1:
            shp = datasets[0]["shape"]
            if len(shp) == 1:
                result["data_type"] = "1d_data"
                result["suggested_agents"] = [0]
                result["primary_suggestion"] = 0
            elif len(shp) == 2:
                result["data_type"] = "image"
                result["suggested_agents"] = [1]
                result["primary_suggestion"] = 1
            elif len(shp) == 3:
                result["data_type"] = "hyperspectral"
                result["suggested_agents"] = [2]
                result["primary_suggestion"] = 2
            else:
                result["data_type"] = "nd_data"
                result["suggested_agents"] = []
        else:
            result["data_type"] = "unknown"
            result["suggested_agents"] = []

    def _register_all_tools(self):
        """Register all tools with OpenAI format."""
        
        # =====================================================================
        # 1. EXAMINE DATA
        # =====================================================================
        def examine_data(data_path: str) -> str:
            """
            Examine a data file or directory to determine its type and
            characteristics.  Supports single files and directories
            containing multiple spectra.

            For directories, JSON files whose stems match a data file
            (e.g. ``spec_5K.json`` ↔ ``spec_5K.csv``) are reported as
            per-file sidecar metadata in ``sidecar_json_files``, separate
            from global ``metadata_files``.
            """
            print(f"  ⚡ Tool: Examining data at {data_path}...")
            
            path = Path(data_path)
            if not path.exists():
                return json.dumps({
                    "status": "error",
                    "message": f"File not found: {data_path}"
                })
            
            result = {
                "status": "success",
                "path": str(path.absolute()),
            }
            
            try:
                # ============================================================
                # DIRECTORY: Multiple files (series)
                # ============================================================
                if path.is_dir():
                    files = list(path.iterdir())
                    files = [f for f in files if f.is_file() and not f.name.startswith('.')]
                    
                    result["is_directory"] = True
                    result["file_count"] = len(files)
                    
                    if not files:
                        result["status"] = "error"
                        result["message"] = "Directory is empty"
                        return json.dumps(result)
                    
                    # Look for metadata files — distinguish sidecar JSONs from
                    # global metadata so the orchestrator knows sidecars exist.
                    non_json_stems = {
                        f.stem for f in files if f.suffix.lower() != ".json"
                    }
                    sidecar_jsons = [
                        f for f in files
                        if f.suffix.lower() == ".json"
                        and f.name.lower() not in _GLOBAL_METADATA_NAMES
                        and f.stem in non_json_stems
                    ]
                    global_meta_files = [
                        f for f in files
                        if (
                            (f.suffix.lower() == ".json" and f not in sidecar_jsons)
                            or "metadata" in f.name.lower()
                            or f.name.lower() in ["info.txt", "description.txt", "readme.txt"]
                        )
                    ]

                    if global_meta_files:
                        result["metadata_files"] = [f.name for f in global_meta_files]
                        result["metadata_hint"] = (
                            f"Found potential metadata file(s): "
                            f"{[f.name for f in global_meta_files]}"
                        )
                    if sidecar_jsons:
                        result["sidecar_json_files"] = [f.name for f in sidecar_jsons]
                        result["sidecar_hint"] = (
                            f"Found {len(sidecar_jsons)} per-file JSON sidecar(s) "
                            f"(may contain series variable values)"
                        )

                    # Get data file extensions (excluding metadata and sidecars)
                    excluded = set(global_meta_files) | set(sidecar_jsons)
                    data_files = [f for f in files if f not in excluded]
                    extensions = set(f.suffix.lower() for f in data_files)
                    result["extensions"] = list(extensions)
                    
                    # Categorize by extension
                    csv_files = [f for f in data_files if f.suffix.lower() in ['.csv', '.txt', '.tsv'] 
                                 and 'metadata' not in f.name.lower() 
                                 and f.name.lower() not in ['info.txt', 'description.txt', 'readme.txt']]
                    npy_files = [f for f in data_files if f.suffix.lower() == '.npy']
                    image_files = [f for f in data_files if f.suffix.lower() in ['.tif', '.tiff', '.png', '.jpg', '.jpeg', '.bmp']]
                    
                    if csv_files:
                        result["data_type"] = "tabular_series"
                        result["series_count"] = len(csv_files)
                        result["suggested_agents"] = [0]  # CurveFitting
                        result["primary_suggestion"] = 0
                        result["data_files"] = sorted([f.name for f in csv_files[:10]])
                        if len(csv_files) > 10:
                            result["data_files"].append(f"... and {len(csv_files) - 10} more")
                        result["note"] = f"Directory contains {len(csv_files)} tabular data files (CSV/TXT) - curves, spectra, time series, etc."
                        
                    elif npy_files:
                        # Check first NPY to determine type
                        first_npy = np.load(str(npy_files[0]))
                        if first_npy.ndim == 1:
                            result["data_type"] = "tabular_series"
                            result["suggested_agents"] = [0]
                            result["primary_suggestion"] = 0
                            result["note"] = f"Directory contains {len(npy_files)} NPY files (1D data)"
                        elif first_npy.ndim == 2:
                            # Distinguish images from tabular data
                            is_image = (
                                min(first_npy.shape) >= 64
                                and max(first_npy.shape) / min(first_npy.shape) <= 4
                            )
                            if is_image:
                                result["data_type"] = "image_series"
                                result["suggested_agents"] = [1]  # ImageAnalysis
                                result["primary_suggestion"] = 1
                                result["note"] = (
                                    f"Directory contains {len(npy_files)} NPY files "
                                    f"({first_npy.shape[0]}x{first_npy.shape[1]}, "
                                    f"{first_npy.dtype}) — detected as image series"
                                )
                            else:
                                result["data_type"] = "tabular_series"
                                result["suggested_agents"] = [0]
                                result["primary_suggestion"] = 0
                                result["note"] = f"Directory contains {len(npy_files)} NPY files (2D tabular data)"
                        else:
                            result["data_type"] = "hyperspectral_series"
                            result["suggested_agents"] = [2]
                            result["primary_suggestion"] = 2
                            result["note"] = f"Directory contains {len(npy_files)} NPY files (3D datacubes)"
                        
                        result["series_count"] = len(npy_files)
                        result["data_files"] = sorted([f.name for f in npy_files[:10]])
                        if len(npy_files) > 10:
                            result["data_files"].append(f"... and {len(npy_files) - 10} more")
                        
                    elif image_files:
                        result["data_type"] = "image_series"
                        result["series_count"] = len(image_files)
                        result["suggested_agents"] = [1]  # ImageAnalysis
                        result["primary_suggestion"] = 1
                        result["data_files"] = sorted([f.name for f in image_files[:10]])
                        if len(image_files) > 10:
                            result["data_files"].append(f"... and {len(image_files) - 10} more")
                        result["note"] = f"Directory contains {len(image_files)} image files - microscopy, photos, etc."
                    
                    else:
                        result["data_type"] = "unknown"
                        result["message"] = f"Directory contains unsupported file types: {extensions}"
                    
                    # Store in orchestrator state
                    self.orch.current_data_path = str(path.absolute())
                    self.orch.current_data_type = result.get("data_type")
                    
                    return json.dumps(result)
                
                # ============================================================
                # SINGLE FILE
                # ============================================================
                file_size = path.stat().st_size
                extension = path.suffix.lower()
                
                result["is_directory"] = False
                result["file_name"] = path.name
                result["file_size_bytes"] = file_size
                result["extension"] = extension
                
                # Determine data type based on extension and content
                if extension in ['.tif', '.tiff', '.png', '.jpg', '.jpeg', '.bmp']:
                    result["data_type"] = "microscopy"
                    result["suggested_agents"] = [1]  # ImageAnalysis
                    
                    # Try to load and get shape
                    try:
                        from ...skills._shared.image_processor import load_image
                        img = load_image(str(path))
                        result["shape"] = list(img.shape)
                        result["dtype"] = str(img.dtype)
                        
                        # Suggest based on image characteristics
                        if len(img.shape) == 2:
                            h, w = img.shape
                        else:
                            h, w = img.shape[:2]
                        
                        result["image_size"] = f"{w}x{h}"
                        result["primary_suggestion"] = 1  # ImageAnalysis
                            
                    except Exception as e:
                        result["load_error"] = str(e)
                
                elif extension == '.npy':
                    # Could be 1D data, 2D data/image, series, or hyperspectral
                    data = np.load(str(path))
                    result["shape"] = list(data.shape)
                    result["dtype"] = str(data.dtype)
                    
                    if data.ndim == 1:
                        result["data_type"] = "1d_data"
                        result["suggested_agents"] = [0]  # CurveFitting
                        result["primary_suggestion"] = 0
                        result["n_points"] = data.shape[0]
                        result["note"] = "Single 1D array - curve, spectrum, time series, etc."
                        
                    elif data.ndim == 2:
                        # Check if it's a series (N x points) or single data (points x 2) or image
                        if data.shape[1] == 2:
                            # Single data with x,y columns
                            result["data_type"] = "1d_data"
                            result["suggested_agents"] = [0]
                            result["primary_suggestion"] = 0
                            result["n_points"] = data.shape[0]
                            result["note"] = "Single dataset with (x, y) columns"
                        elif data.shape[0] == 2:
                            # Single data with x,y rows
                            result["data_type"] = "1d_data"
                            result["suggested_agents"] = [0]
                            result["primary_suggestion"] = 0
                            result["n_points"] = data.shape[1]
                            result["note"] = "Single dataset with (x, y) rows"
                        elif data.shape[0] > 2 and data.shape[1] > 2:
                            # Could be series of 1D data OR 2D image
                            # Heuristic: if one dimension is much smaller, likely a series
                            if data.shape[0] < 100 and data.shape[1] > 100:
                                # Likely N datasets of M points each
                                result["data_type"] = "1d_series"
                                result["suggested_agents"] = [0]
                                result["primary_suggestion"] = 0
                                result["series_count"] = data.shape[0]
                                result["n_points"] = data.shape[1]
                                result["note"] = f"Series of {data.shape[0]} datasets, each with {data.shape[1]} points"
                            elif data.shape[1] < 100 and data.shape[0] > 100:
                                # Likely M points x N datasets (transposed)
                                result["data_type"] = "1d_series"
                                result["suggested_agents"] = [0]
                                result["primary_suggestion"] = 0
                                result["series_count"] = data.shape[1]
                                result["n_points"] = data.shape[0]
                                result["note"] = f"Series of {data.shape[1]} datasets, each with {data.shape[0]} points (may need transpose)"
                            elif (
                                min(data.shape) >= 64
                                and max(data.shape) / min(data.shape) <= 4
                                and data.dtype in (
                                    np.uint8, np.uint16, np.int16,
                                    np.float32, np.float64,
                                )
                            ):
                                # Large, roughly-square array — almost certainly an image
                                result["data_type"] = "image"
                                result["suggested_agents"] = [1]
                                result["primary_suggestion"] = 1
                                result["note"] = (
                                    f"2D array ({data.shape[0]}x{data.shape[1]}, "
                                    f"{data.dtype}) — detected as image"
                                )
                            else:
                                # Ambiguous - could be image or data matrix
                                # Try to infer from metadata if available
                                result["data_type"] = "2d_data_ambiguous"
                                result["suggested_agents"] = [0, 1]
                                result["primary_suggestion"] = None  # No clear suggestion
                                result["note"] = (
                                    f"Ambiguous 2D array ({data.shape[0]}x{data.shape[1]}). Could be:\n"
                                    f"  - Microscopy image → Agent 1 (ImageAnalysisAgent)\n"
                                    f"  - Series of 1D data (rows or columns) → Agent 0 (CurveFittingAgent)\n"
                                    f"  - 2D spectral slice → Agent 2 (HyperspectralAnalysisAgent)\n"
                                    f"Check metadata or ask user to clarify."
                                )
                                result["disambiguation_needed"] = True
                                result["disambiguation_questions"] = [
                                    "Is this a microscopy/image?",
                                    "Is this a matrix where each row (or column) is a separate spectrum/curve?",
                                    "What technique was used to acquire this data?"
                                ]
                            
                    elif data.ndim == 3:
                        n_channels = data.shape[2]
                        if n_channels in (2, 3, 4):
                            # Few channels: multi-channel image (2-ch AFM, RGB, RGBA)
                            result["data_type"] = "image"
                            result["suggested_agents"] = [1]  # ImageAnalysis
                            result["primary_suggestion"] = 1
                            result["channels"] = n_channels
                            result["note"] = (
                                f"3D array ({data.shape[0]}x{data.shape[1]}, "
                                f"{n_channels} channels, {data.dtype}) — "
                                f"detected as multi-channel image"
                            )
                        else:
                            # Many channels: spectral datacube
                            result["data_type"] = "hyperspectral"
                            result["suggested_agents"] = [2]  # Hyperspectral
                            result["primary_suggestion"] = 2
                            result["spatial_shape"] = list(data.shape[:2])
                            result["spectral_channels"] = n_channels
                            result["note"] = (
                                f"3D datacube: {data.shape[0]}x{data.shape[1]} "
                                f"spatial, {n_channels} channels"
                            )
                    
                    else:
                        result["data_type"] = "nd_data"
                        result["note"] = f"{data.ndim}D array - may need custom handling"
                        result["suggested_agents"] = []
                
                elif extension in ['.csv', '.txt', '.tsv']:
                    result["data_type"] = "tabular"
                    result["suggested_agents"] = [0]  # CurveFitting
                    result["primary_suggestion"] = 0

                    # Try to peek at the file and count rows
                    try:
                        import csv
                        with open(path, 'r') as f:
                            # Read first few lines for preview
                            first_lines = [f.readline().strip() for _ in range(5)]
                            result["preview"] = first_lines

                            # Count total lines (approximate row count)
                            f.seek(0)
                            row_count = sum(1 for _ in f) - 1  # Subtract header
                            result["n_points"] = row_count
                            result["note"] = f"Tabular data with ~{row_count} data points"
                    except Exception as e:
                        result["preview_error"] = str(e)

                elif extension in ['.h5', '.hdf5']:
                    # NeXus / SID-style HDF5 — produced by the SciFiReaders
                    # MCP server (read_scifireaders_file) or any sidpy
                    # pipeline. Surface shape, dimensions, and metadata so
                    # the LLM has enough to route to the right agent without
                    # needing the raw array inline.
                    self._examine_hdf5(path, result)

                else:
                    result["data_type"] = "unknown"
                    result["message"] = f"Unknown file extension: {extension}"
                    result["suggested_agents"] = []
                    # Hint when the user uploaded a vendor format that needs
                    # to be converted via the SciFiReaders MCP server before
                    # SciLink can examine it.
                    from ...ui.config import VENDOR_DATA_EXTENSIONS
                    if extension in VENDOR_DATA_EXTENSIONS:
                        result["hint"] = (
                            f"'{extension}' is a vendor format. If the "
                            "SciFiReaders MCP server is connected, call "
                            "read_scifireaders_file(file_path) to convert "
                            "to NeXus HDF5, then re-run examine_data on "
                            "the resulting '.nxs.h5' file to get shape, "
                            "dimensions, and metadata."
                        )
                
                # Store in orchestrator state
                self.orch.current_data_path = str(path.absolute())
                self.orch.current_data_type = result.get("data_type")
                
                return json.dumps(result)
                
            except Exception as e:
                self.logger.error(f"Error examining data: {e}", exc_info=True)
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=examine_data,
            name="examine_data",
            description=(
                "Examine a data file to determine its type and characteristics. "
                "Returns data type, shape, and suggested analysis agents. "
                "For directories, also detects per-file JSON sidecar metadata "
                "(stem-matched to data files) and reports them separately from "
                "global metadata files."
            ),
            parameters={
                "data_path": {
                    "type": "string",
                    "description": "Path to the data file to examine"
                }
            },
            required=["data_path"]
        )
        
        # =====================================================================
        # 2. CONVERT METADATA
        # =====================================================================
        def convert_metadata(
            text_input: str = None,
            text_file_path: str = None
        ) -> str:
            """
            Convert natural language description to structured metadata JSON.
            """
            print(f"  ⚡ Tool: Converting metadata...")
            
            if text_file_path:
                path = Path(text_file_path)
                if not path.exists():
                    return json.dumps({
                        "status": "error",
                        "message": f"File not found: {text_file_path}"
                    })
                
                # Use the metadata converter
                try:
                    metadata = generate_metadata_json_from_text(
                        input_text_filepath=str(path),
                        api_key=self.orch.api_key,
                        model_name=self.orch.model_name,
                        base_url=self.orch.base_url
                    )
                    
                    if metadata:
                        cpi_warning = self._replace_metadata(metadata)
                        output_path = self.orch.base_dir / "metadata.json"
                        with open(output_path, 'w') as f:
                            json.dump(metadata, f, indent=2)

                        result = {
                            "status": "success",
                            "metadata": metadata,
                            "saved_to": str(output_path)
                        }
                        if cpi_warning:
                            result.update(cpi_warning)
                        return json.dumps(result)
                    else:
                        return json.dumps({
                            "status": "error",
                            "message": "Failed to convert metadata"
                        })

                except Exception as e:
                    self.logger.error(f"Metadata conversion error: {e}", exc_info=True)
                    return json.dumps({
                        "status": "error",
                        "message": str(e)
                    })

            elif text_input:
                # Create temporary file and convert
                temp_path = self.orch.base_dir / "temp_metadata_input.txt"
                with open(temp_path, 'w') as f:
                    f.write(text_input)
                
                try:
                    metadata = generate_metadata_json_from_text(
                        input_text_filepath=str(temp_path),
                        api_key=self.orch.api_key,
                        model_name=self.orch.model_name,
                        base_url=self.orch.base_url
                    )
                    
                    # Clean up temp file
                    temp_path.unlink()
                    
                    if metadata:
                        cpi_warning = self._replace_metadata(metadata)
                        output_path = self.orch.base_dir / "metadata.json"
                        with open(output_path, 'w') as f:
                            json.dump(metadata, f, indent=2)

                        result = {
                            "status": "success",
                            "metadata": metadata,
                            "saved_to": str(output_path)
                        }
                        if cpi_warning:
                            result.update(cpi_warning)
                        return json.dumps(result)
                    else:
                        return json.dumps({
                            "status": "error",
                            "message": "Failed to convert metadata"
                        })

                except Exception as e:
                    if temp_path.exists():
                        temp_path.unlink()
                    return json.dumps({
                        "status": "error",
                        "message": str(e)
                    })

            else:
                return json.dumps({
                    "status": "error",
                    "message": "Must provide either text_input or text_file_path"
                })
        
        self._register_tool(
            func=convert_metadata,
            name="convert_metadata",
            description=(
                "Convert natural language description to structured metadata JSON. "
                "Accepts either direct text input or a path to a text file. "
                "Use this when user provides experimental description in plain text."
            ),
            parameters={
                "text_input": {
                    "type": "string",
                    "description": "Direct text description of the experiment (alternative to file)"
                },
                "text_file_path": {
                    "type": "string",
                    "description": "Path to a .txt file containing experiment description"
                }
            },
            required=[]
        )
        
        # =====================================================================
        # 3. LOAD METADATA
        # =====================================================================
        def load_metadata(json_path: str) -> str:
            """
            Load existing JSON metadata file.

            Can accept either a direct path to a JSON file, or a directory
            path (will search for metadata.json or similar files in the
            directory).  Per-file sidecar JSONs (whose stem matches a data
            file) are excluded from the search so they are not mistakenly
            loaded as global metadata.
            """
            print(f"  ⚡ Tool: Loading metadata from {json_path}...")
            
            path = Path(json_path)
            if not path.exists():
                return json.dumps({
                    "status": "error",
                    "message": f"File/directory not found: {json_path}"
                })
            
            # If directory, search for metadata file
            if path.is_dir():
                # Look for common metadata file names
                metadata_candidates = [
                    path / "metadata.json",
                    path / "meta.json",
                    path / "info.json",
                    path / "experiment.json",
                ]
                
                # Also look for any .json file, but exclude sidecar JSONs
                # (files whose stem matches a data file, e.g. spec_5K.json ↔ spec_5K.csv)
                json_files = list(path.glob("*.json"))
                _data_exts = {
                    ".csv", ".txt", ".tsv", ".xlsx",
                    ".npy", ".tif", ".tiff", ".png", ".jpg", ".jpeg",
                }
                _data_stems = {
                    f.stem
                    for f in path.iterdir()
                    if f.is_file() and f.suffix.lower() in _data_exts
                }
                non_sidecar_jsons = [
                    jf for jf in json_files if jf.stem not in _data_stems
                ]

                # Find the first existing metadata file
                metadata_path = None
                for candidate in metadata_candidates:
                    if candidate.exists():
                        metadata_path = candidate
                        break

                # If no standard name found, use first non-sidecar .json file
                if metadata_path is None and non_sidecar_jsons:
                    metadata_path = non_sidecar_jsons[0]
                
                if metadata_path is None:
                    # ---------------------------------------------------------
                    # Synthesize global metadata from sidecar JSONs
                    # ---------------------------------------------------------
                    # When there is no dedicated metadata file but per-file
                    # sidecars exist, extract fields that are identical across
                    # ALL sidecars as shared (global) metadata.  This lets
                    # users skip writing a separate metadata.json when the
                    # sidecars already contain experiment/sample information.
                    sidecar_paths = [
                        jf for jf in json_files if jf.stem in _data_stems
                    ]
                    if sidecar_paths:
                        try:
                            all_sidecar_dicts = []
                            for sp in sidecar_paths:
                                with open(sp, "r") as _f:
                                    all_sidecar_dicts.append(json.load(_f))

                            if all_sidecar_dicts:
                                # Collect keys shared by every sidecar
                                shared_keys = set(all_sidecar_dicts[0].keys())
                                for sd in all_sidecar_dicts[1:]:
                                    shared_keys &= sd.keys()

                                # Keep only fields whose value is the same in
                                # every sidecar (these describe the experiment,
                                # not the varying control variable).
                                synthesized: dict = {}
                                for key in shared_keys:
                                    values = [sd[key] for sd in all_sidecar_dicts]
                                    ref = values[0]
                                    if all(v == ref for v in values):
                                        synthesized[key] = ref

                                if synthesized:
                                    # Normalize to canonical schema
                                    is_conformant, _ = check_schema_conformance(synthesized)
                                    if not is_conformant:
                                        normed, _ = normalize_metadata_dict(synthesized)
                                        re_ok, _ = check_schema_conformance(normed)
                                        if not re_ok:
                                            try:
                                                llm_result = normalize_metadata_dict_with_llm(
                                                    synthesized, self.orch.model, self.logger
                                                )
                                                if llm_result:
                                                    for k, v in synthesized.items():
                                                        if k not in llm_result:
                                                            llm_result[k] = v
                                                    synthesized = llm_result
                                            except Exception:
                                                synthesized = normed
                                        else:
                                            synthesized = normed

                                    cpi_warning = self._replace_metadata(synthesized)
                                    output_path = self.orch.base_dir / "metadata.json"
                                    with open(output_path, 'w') as f:
                                        json.dump(synthesized, f, indent=2)
                                    print(
                                        f"    Synthesized global metadata from "
                                        f"{len(sidecar_paths)} sidecar JSON(s)"
                                    )

                                    required_fields = ["experiment_type", "experiment", "sample"]
                                    missing = [f for f in required_fields if f not in synthesized]
                                    status = "warning" if missing else "success"
                                    result_payload = {
                                        "status": status,
                                        "source": "synthesized_from_sidecars",
                                        "num_sidecars": len(sidecar_paths),
                                        "metadata": synthesized,
                                        "experiment_type": synthesized.get("experiment_type"),
                                        "technique": (
                                            synthesized.get("experiment", {}).get("technique")
                                            if isinstance(synthesized.get("experiment"), dict)
                                            else synthesized.get("technique")
                                        ),
                                        "material": (
                                            synthesized.get("sample", {}).get("material")
                                            if isinstance(synthesized.get("sample"), dict)
                                            else synthesized.get("material")
                                        ),
                                    }
                                    if missing:
                                        result_payload["message"] = (
                                            f"Metadata synthesized from sidecar JSONs "
                                            f"but missing recommended fields: {missing}"
                                        )
                                    if cpi_warning:
                                        result_payload.update(cpi_warning)
                                    return json.dumps(result_payload)
                        except Exception as e:
                            self.logger.warning(
                                f"Failed to synthesize metadata from sidecars: {e}"
                            )

                    # Look for .txt description files
                    txt_candidates = [
                        path / "metadata.txt",
                        path / "description.txt",
                        path / "info.txt",
                    ]
                    for candidate in txt_candidates:
                        if candidate.exists():
                            return json.dumps({
                                "status": "info",
                                "message": f"Found text description file: {candidate.name}. Use convert_metadata to convert it to JSON.",
                                "text_file": str(candidate)
                            })

                    return json.dumps({
                        "status": "error",
                        "message": f"No metadata file found in directory: {json_path}"
                    })
                
                path = metadata_path
                print(f"    Found metadata file: {path.name}")
            
            try:
                with open(path, 'r') as f:
                    metadata = json.load(f)
                
                # Normalize metadata to canonical schema if needed
                is_conformant, issues = check_schema_conformance(metadata)
                if not is_conformant:
                    normalized, was_modified = normalize_metadata_dict(metadata)  # Tier 1
                    re_ok, _ = check_schema_conformance(normalized)
                    if not re_ok:
                        # Tier 2: LLM normalization for remaining gaps
                        try:
                            llm_result = normalize_metadata_dict_with_llm(
                                metadata, self.orch.model, self.logger
                            )
                            if llm_result:
                                # Preserve non-schema keys from the original
                                for k, v in metadata.items():
                                    if k not in llm_result:
                                        llm_result[k] = v
                                metadata = llm_result
                        except Exception as e:
                            self.logger.warning(f"LLM metadata normalization failed: {e}")
                            if was_modified:
                                metadata = normalized
                    else:
                        metadata = normalized

                # Always store metadata (possibly normalized)
                cpi_warning = self._replace_metadata(metadata)

                # Validate basic structure
                required_fields = ["experiment_type", "experiment", "sample"]
                missing = [f for f in required_fields if f not in metadata]

                if missing:
                    result = {
                        "status": "warning",
                        "message": f"Metadata loaded but missing recommended fields: {missing}",
                        "metadata_file": path.name,
                        "metadata": metadata,
                        "experiment_type": metadata.get("experiment_type"),
                        "technique": metadata.get("experiment", {}).get("technique") if isinstance(metadata.get("experiment"), dict) else metadata.get("technique"),
                        "material": metadata.get("sample", {}).get("material") if isinstance(metadata.get("sample"), dict) else metadata.get("material")
                    }
                    if cpi_warning:
                        result.update(cpi_warning)
                    return json.dumps(result)

                result = {
                    "status": "success",
                    "metadata_file": path.name,
                    "metadata": metadata,
                    "experiment_type": metadata.get("experiment_type"),
                    "technique": metadata.get("experiment", {}).get("technique"),
                    "material": metadata.get("sample", {}).get("material")
                }
                if cpi_warning:
                    result.update(cpi_warning)
                return json.dumps(result)
                
            except json.JSONDecodeError as e:
                return json.dumps({
                    "status": "error",
                    "message": f"Invalid JSON: {e}"
                })
            except Exception as e:
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=load_metadata,
            name="load_metadata",
            description=(
                "Load experiment metadata. "
                "Can accept a direct path to a .json file OR a directory path. "
                "When given a directory it will: (1) look for a dedicated metadata file "
                "(metadata.json, meta.json, info.json, etc.), or (2) if none exists, "
                "automatically synthesize global metadata from per-file sidecar JSONs "
                "by extracting fields that are shared across all sidecars and "
                "normalizing them into the canonical schema. "
                "Use this for any directory containing metadata — whether as a "
                "single file or as per-file sidecars."
            ),
            parameters={
                "json_path": {
                    "type": "string",
                    "description": "Path to JSON metadata file OR directory containing metadata"
                }
            },
            required=["json_path"]
        )
        
        # =====================================================================
        # 4. SELECT AGENT
        # =====================================================================
        def select_agent(
            agent_id: int,
            reasoning: str = None
        ) -> str:
            """
            Set the selected analysis agent. The chat LLM decides which agent to use
            based on data type, metadata, and image preview (if applicable).
            
            Agent IDs:
                0: CurveFittingAgent - 1D curves, spectra
                1: ImageAnalysisAgent - all image types
                2: HyperspectralAnalysisAgent - spectral datacubes
            """
            print(f"  ⚡ Tool: Setting agent to {agent_id}...")
            
            if agent_id not in self.AGENT_NAMES:
                return json.dumps({
                    "status": "error",
                    "message": f"Invalid agent_id: {agent_id}. Valid IDs: {list(self.AGENT_NAMES.keys())}"
                })
            
            self.orch.selected_agent_id = agent_id
            
            return json.dumps({
                "status": "success",
                "agent_id": agent_id,
                "agent_name": self.AGENT_NAMES.get(agent_id),
                "description": self.AGENT_DESCRIPTIONS.get(agent_id),
                "reasoning": reasoning or "Selected by user/LLM"
            })
        
        self._register_tool(
            func=select_agent,
            name="select_agent",
            description=(
                "Set the analysis agent to use. Call this after examining data and metadata. "
                "Agent IDs: 0=CurveFitting (1D data), 1=ImageAnalysis (all images), 2=Hyperspectral (3D datacubes)"
            ),
            parameters={
                "agent_id": {
                    "type": "integer",
                    "description": "Agent ID to use (0=CurveFitting, 1=ImageAnalysis, 2=Hyperspectral)"
                },
                "reasoning": {
                    "type": "string",
                    "description": "Brief explanation of why this agent was chosen"
                }
            },
            required=["agent_id"]
        )
        
        # =====================================================================
        # 4b. PREVIEW IMAGE (for microscopy agent selection)
        # =====================================================================
        def preview_image(image_path: str = None) -> str:
            """
            Load and return a preview of a microscopy image for the LLM to analyze.
            Use this to visually inspect the image before analysis.
            """
            print(f"  ⚡ Tool: Loading image preview...")
            
            if image_path is None:
                image_path = self.orch.current_data_path
            
            if image_path is None:
                return json.dumps({
                    "status": "error",
                    "message": "No image path provided. Use examine_data first."
                })
            
            path = Path(image_path)
            if not path.exists():
                return json.dumps({
                    "status": "error", 
                    "message": f"File not found: {image_path}"
                })
            
            # Accept standard image formats, plus the array-container formats
            # load_image can read (.npy/.h5/.hdf5). The latter are generic
            # containers — a .npy may hold a spectrum or datacube, not an
            # image — so the loaded array's shape is validated below.
            image_extensions = ['.tif', '.tiff', '.png', '.jpg', '.jpeg', '.bmp']
            array_extensions = ['.npy', '.h5', '.hdf5']
            suffix = path.suffix.lower()
            if suffix not in image_extensions and suffix not in array_extensions:
                return json.dumps({
                    "status": "error",
                    "message": (
                        f"Unsupported file type: {path.suffix}. preview_image "
                        "accepts .tif/.tiff/.png/.jpg/.jpeg/.bmp, or "
                        ".npy/.h5/.hdf5 holding a 2-D or RGB image array."
                    )
                })

            try:
                from ...skills._shared.image_processor import load_image
                import base64
                from io import BytesIO
                from PIL import Image

                # load_image handles .npy/.h5 too, and normalizes them to uint8
                img_array = load_image(str(path))

                # An array container may hold a non-image. Only a 2-D array,
                # or 3-D with 3/4 colour channels, is a previewable image.
                if img_array.ndim == 3 and img_array.shape[-1] == 1:
                    img_array = img_array[:, :, 0]   # singleton channel -> 2-D
                is_image = (
                    img_array.ndim == 2
                    or (img_array.ndim == 3 and img_array.shape[-1] in (3, 4))
                )
                if not is_image:
                    if img_array.ndim == 1:
                        guess = "a 1-D signal / spectrum"
                    elif img_array.ndim == 3:
                        guess = (f"a {img_array.shape[-1]}-channel datacube "
                                 "(e.g. hyperspectral)")
                    else:
                        guess = f"a {img_array.ndim}-D array"
                    return json.dumps({
                        "status": "error",
                        "message": (
                            f"{path.name} holds an array of shape "
                            f"{list(img_array.shape)} — not a previewable "
                            f"image; it looks like {guess}. preview_image "
                            "renders 2-D or RGB images only."
                        )
                    })

                # Get basic stats
                shape = img_array.shape
                dtype = str(img_array.dtype)

                # Convert to PIL for resizing and encoding
                pil_img = Image.fromarray(img_array)

                # Resize for preview (max 512px)
                max_dim = 512
                if max(pil_img.size) > max_dim:
                    ratio = max_dim / max(pil_img.size)
                    new_size = (int(pil_img.size[0] * ratio), int(pil_img.size[1] * ratio))
                    pil_img = pil_img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Convert to base64
                buffer = BytesIO()
                pil_img.convert('RGB').save(buffer, format='JPEG', quality=85)
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                
                return json.dumps({
                    "status": "success",
                    "image_path": str(path),
                    "shape": list(shape),
                    "dtype": dtype,
                    "preview_size": list(pil_img.size),
                    "image_base64": img_base64,
                    "guidance": (
                        "Examine this image. For any microscopy/image data, use "
                        "ImageAnalysisAgent (ID: 1)."
                    )
                })
                
            except Exception as e:
                self.logger.error(f"Image preview error: {e}", exc_info=True)
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=preview_image,
            name="preview_image",
            description=(
                "Load a microscopy image preview for visual inspection. "
                "Accepts .png/.tif/.jpg/.bmp and .npy/.h5 array files that "
                "hold a 2-D or RGB image (a non-image array — spectrum, "
                "datacube — is rejected with an explanation). Returns the "
                "image as base64 for you to examine."
            ),
            parameters={
                "image_path": {
                    "type": "string",
                    "description": "Path to image file (uses current data path if not specified)"
                }
            },
            required=[]
        )
        
        # =====================================================================
        # 5. RUN ANALYSIS
        # =====================================================================
        def run_analysis(
            data_path: str = None,
            agent_id: int = None,
            analysis_goal: str = None,
            objective: str = None,
            hints: str = None,
            auxiliary_data: str = None,
            auxiliary_label: str = None,
            skill = None,  # str | list[str] | None (PR 3 multi-skill)
            series_metadata: str = None,
            task_mode: str = None,
            prior_analysis_paths: List[str] = None,
            literature_file: str = None,
        ) -> str:
            """
            Execute analysis with the selected or specified agent.

            Each analysis run creates a unique output directory under results/
            to ensure traceability and prevent output collisions when analyzing
            multiple datasets with the same agent.

            For agents that execute LLM-generated code (CurveFitting, Hyperspectral),
            a sandbox check is performed. If no sandbox is detected, the user is
            prompted to confirm before proceeding.

            auxiliary_data and auxiliary_label can provide a complementary dataset
            (e.g. TGA curve alongside DSC, or a microscopy image) as context for
            the analysis without fitting/unmixing it. Supported by CurveFitting
            and Hyperspectral agents.

            Series metadata resolution (in priority order):

            1. **Explicit ``series_metadata`` parameter** — a JSON string
               describing the independent variable.  When ``values`` is a
               dict mapping filenames to values, files are automatically
               sorted by value for correct physical ordering and the dict
               is converted to a sorted list before passing to the agent.
               Expected format::

                   {"variable": "temperature",
                    "values": {"spec_5K.csv": 5, ...}, "unit": "K"}

            2. **Per-file JSON sidecars** — if the data directory contains
               JSON files whose stems match data files (e.g.
               ``spec_5K.json`` ↔ ``spec_5K.csv``), the system loads them,
               identifies numeric fields that vary across files, and uses
               LLM reasoning (with experimental context from *objective*,
               *hints*, and loaded metadata) to decide which field, if any,
               is the true control variable.  Full sidecar contents are
               stored in ``current_metadata["per_file_metadata"]`` for
               downstream agent access.

            3. **User prompt** — if neither of the above yields series
               metadata, returns a ``needs_series_metadata`` status.  When
               sidecars were loaded but the LLM could not identify a
               control variable, the sidecar contents are included in the
               response so the orchestrator can show them to the user.

            Output directory format: results/analysis_{dataset_name}_{timestamp}_{counter}/
            """
            print(f"  ⚡ Tool: Running analysis...")
            
            # Use current state if not provided
            if data_path is None:
                data_path = self.orch.current_data_path
            
            if agent_id is None:
                agent_id = self.orch.selected_agent_id
            
            # Validate inputs
            if data_path is None:
                return json.dumps({
                    "status": "error",
                    "message": "No data path provided. Use examine_data first."
                })
            
            if agent_id is None:
                return json.dumps({
                    "status": "error",
                    "message": "No agent selected. Use select_agent first."
                })
            
            if self.orch.current_metadata is None:
                # When a data directory contains sidecar JSONs, allow
                # run_analysis to proceed so the sidecar extraction code
                # below can populate metadata automatically.
                data_p = Path(data_path)
                has_sidecars = False
                if data_p.is_dir():
                    _all = [f for f in data_p.iterdir() if f.is_file() and not f.name.startswith('.')]
                    _data = [f for f in _all if f.suffix.lower() != ".json"]
                    _smap, _ = _detect_sidecar_jsons(_data, _all)
                    has_sidecars = bool(_smap)
                if has_sidecars:
                    self.orch.current_metadata = {}
                else:
                    return json.dumps({
                        "status": "error",
                        "message": "No metadata available. Use load_metadata or convert_metadata first."
                    })
            
            try:
                # === Handle directory input - filter out metadata files ===
                path = Path(data_path)
                actual_data_input = data_path  # Default: pass as-is
                
                if path.is_dir():
                    # Get all files excluding metadata
                    all_files = [f for f in path.iterdir() if f.is_file() and not f.name.startswith('.')]
                    
                    # Filter out metadata files
                    data_files = []
                    for f in all_files:
                        is_metadata = (
                            f.suffix.lower() == '.json' or
                            'metadata' in f.name.lower() or
                            f.name.lower() in ['info.txt', 'description.txt', 'readme.txt', 'readme.md']
                        )
                        if not is_metadata:
                            data_files.append(f)
                    
                    if not data_files:
                        return json.dumps({
                            "status": "error",
                            "message": "No data files found in directory (only metadata files present)"
                        })
                    
                    # Sort for consistent ordering
                    data_files = sorted(data_files, key=lambda x: x.name)
                    
                    print(f"    Found {len(data_files)} data files (excluded metadata)")
                    
                    # Pass as list of file paths for series analysis
                    actual_data_input = [str(f) for f in data_files]
                    
                    # If only one file, pass as string (single spectrum mode)
                    if len(actual_data_input) == 1:
                        actual_data_input = actual_data_input[0]
                        print(f"    Single file in directory, using single spectrum mode")
                    else:
                        print(f"    Series mode: passing {len(actual_data_input)} files")
                        for i, fp in enumerate(actual_data_input[:3]):
                            print(f"      [{i}] {Path(fp).name}")
                        if len(actual_data_input) > 3:
                            print(f"      ... and {len(actual_data_input) - 3} more")
                
                # === Handle series metadata ===
                is_series = isinstance(actual_data_input, list) and len(actual_data_input) > 1
                has_series_meta = (
                    isinstance(self.orch.current_metadata, dict)
                    and "series" in self.orch.current_metadata
                )

                if series_metadata is not None:
                    # Parse and inject series metadata from the tool call
                    try:
                        parsed_series = json.loads(series_metadata) if isinstance(series_metadata, str) else series_metadata
                        self.orch.current_metadata["series"] = parsed_series
                        has_series_meta = True
                    except (json.JSONDecodeError, TypeError) as e:
                        self.logger.warning(f"Failed to parse series_metadata: {e}")

                # === Try to extract series metadata from sidecar JSON files ===
                if is_series and not has_series_meta and path.is_dir():
                    sidecar_map, _global_jsons = _detect_sidecar_jsons(
                        data_files, all_files
                    )
                    if sidecar_map:
                        print(
                            f"    Found {len(sidecar_map)} sidecar JSON file(s) "
                            f"paired with data files"
                        )
                        extracted_series, per_file_meta = (
                            _extract_series_from_sidecars(
                                sidecar_map,
                                data_files,
                                self.logger,
                                model=self.orch.model,
                                experimental_context={
                                    "objective": objective,
                                    "hints": hints,
                                    "metadata": self.orch.current_metadata,
                                },
                            )
                        )
                        # Store per-file metadata for agent access
                        if per_file_meta:
                            self.orch.current_metadata[
                                "per_file_metadata"
                            ] = per_file_meta

                            # Synthesize normalized global metadata from
                            # shared invariant fields across all sidecars,
                            # unless current_metadata already contains
                            # normalized top-level sections.
                            _has_global = any(
                                k in self.orch.current_metadata
                                for k in ("experiment", "sample", "instrument")
                            )
                            if not _has_global:
                                _all_dicts = list(per_file_meta.values())
                                _shared = set(_all_dicts[0].keys())
                                for _sd in _all_dicts[1:]:
                                    _shared &= _sd.keys()
                                _synth: dict = {}
                                for _k in _shared:
                                    _vals = [_sd[_k] for _sd in _all_dicts]
                                    if all(v == _vals[0] for v in _vals):
                                        _synth[_k] = _vals[0]
                                if _synth:
                                    try:
                                        _ok, _ = check_schema_conformance(_synth)
                                        if not _ok:
                                            _normed, _ = normalize_metadata_dict(_synth)
                                            _re_ok, _ = check_schema_conformance(_normed)
                                            if not _re_ok:
                                                _llm = normalize_metadata_dict_with_llm(
                                                    _synth, self.orch.model, self.logger
                                                )
                                                if _llm:
                                                    for _k2, _v2 in _synth.items():
                                                        if _k2 not in _llm:
                                                            _llm[_k2] = _v2
                                                    _synth = _llm
                                                else:
                                                    _synth = _normed
                                            else:
                                                _synth = _normed
                                        # Merge normalized global fields into
                                        # current_metadata without overwriting
                                        # per_file_metadata or series.
                                        for _k3, _v3 in _synth.items():
                                            if _k3 not in self.orch.current_metadata:
                                                self.orch.current_metadata[_k3] = _v3
                                        print(
                                            f"    Synthesized global metadata from "
                                            f"{len(per_file_meta)} sidecar(s)"
                                        )
                                    except Exception as _e:
                                        self.logger.warning(
                                            "Failed to synthesize global metadata "
                                            "in run_analysis: %s", _e
                                        )

                        # Auto-populate series metadata if extraction succeeded
                        if extracted_series is not None:
                            self.orch.current_metadata["series"] = extracted_series
                            has_series_meta = True
                            print(
                                f"    Auto-extracted series variable "
                                f"'{extracted_series['variable']}' from sidecar JSONs"
                            )
                            # In co-pilot / autopilot modes, let the user
                            # know which control variable was extracted and
                            # give them a chance to confirm or correct it
                            # before proceeding with the analysis.
                            mode = self.orch.analysis_mode.value
                            if mode in ("co-pilot", "autopilot"):
                                values = extracted_series.get("values", {})
                                unit = extracted_series.get("unit", "")
                                # Build a readable summary of the mapping
                                sample_items = list(values.items())[:5]
                                mapping_lines = [
                                    f"  {fname}: {val}"
                                    for fname, val in sample_items
                                ]
                                if len(values) > 5:
                                    mapping_lines.append(
                                        f"  ... and {len(values) - 5} more"
                                    )
                                mapping_str = "\n".join(mapping_lines)
                                return json.dumps({
                                    "status": "series_variable_extracted",
                                    "message": (
                                        f"Auto-extracted series control variable "
                                        f"'{extracted_series['variable']}'"
                                        f"{(' (' + unit + ')') if unit else ''} "
                                        f"from per-file sidecar JSON metadata. "
                                        f"File-to-value mapping:\n{mapping_str}\n\n"
                                        f"Present this to the user and ask them "
                                        f"to confirm it is correct before "
                                        f"proceeding.\n"
                                        f"- If the user CONFIRMS: re-call "
                                        f"run_analysis with the same parameters "
                                        f"(no series_metadata needed — it is "
                                        f"already stored).\n"
                                        f"- If the user DISAGREES or wants a "
                                        f"different variable: ask them for the "
                                        f"correct variable name, values, and "
                                        f"unit, then re-call run_analysis with "
                                        f"an explicit series_metadata parameter "
                                        f"containing the corrected mapping. "
                                        f"The explicit parameter will override "
                                        f"the auto-extracted one."
                                    ),
                                    "variable": extracted_series["variable"],
                                    "unit": unit,
                                    "values": values,
                                    "num_files": len(values),
                                })

                if is_series and not has_series_meta:
                    num_files = len(actual_data_input)

                    # If per-file sidecar metadata was loaded, include it in
                    # the prompt so the orchestrator LLM can show the user
                    # what each file already contains.
                    per_file = self.orch.current_metadata.get(
                        "per_file_metadata"
                    )
                    sidecar_note = ""
                    if per_file:
                        sidecar_note = (
                            " Per-file JSON sidecar metadata was found but "
                            "none of the recorded fields could be "
                            "confidently identified as the control variable. "
                            "Show the user the sidecar contents below and "
                            "ask them to confirm which field (if any) is the "
                            "control variable, or to specify it manually."
                        )

                    prompt_payload = {
                        "status": "needs_series_metadata",
                        "message": (
                            f"Detected {num_files} spectra (series mode) but "
                            "no series metadata found. "
                            "Series metadata describes the experimental "
                            "variable that changes across spectra "
                            "(e.g. temperature, concentration, voltage). "
                            "Ask the user what variable changes across the "
                            "spectra, the range or values, and the units. "
                            "The user can describe this naturally — e.g. "
                            "'temperature from 300 to 500 K in 50 K steps' "
                            "or 'concentration: 0.1, 0.2, 0.5 mM'. "
                            "Use the filenames and the user's response to "
                            "build the values dict mapping each filename to "
                            "its value, then re-call run_analysis with the "
                            "series_metadata parameter. "
                            "Files will be sorted by value automatically for "
                            "correct trend analysis."
                            + sidecar_note
                        ),
                        "num_spectra": num_files,
                        "expected_format": {
                            "variable": "<variable name, e.g. temperature>",
                            "values": {
                                "<filename>": "<value>",
                                "...": "...",
                            },
                            "unit": "<unit string, e.g. K, mM, V>",
                        },
                        "files": [
                            Path(f).name for f in actual_data_input
                        ],
                    }
                    if per_file:
                        prompt_payload["per_file_sidecar_metadata"] = per_file

                    return json.dumps(prompt_payload)

                # Sort files by series values for correct physical ordering
                if is_series and has_series_meta:
                    series_info = self.orch.current_metadata.get("series", {})
                    values = series_info.get("values")
                    if isinstance(values, dict):
                        # Map filenames to full paths
                        name_to_path = {Path(f).name: f for f in actual_data_input}
                        # Build sorted (path, value) pairs by value
                        paired = []
                        for fname, val in values.items():
                            full_path = name_to_path.get(fname)
                            if full_path is not None:
                                try:
                                    paired.append((full_path, float(val)))
                                except (TypeError, ValueError):
                                    paired.append((full_path, val))
                        # Sort by value (numeric sort when possible)
                        try:
                            paired.sort(key=lambda x: x[1])
                        except TypeError:
                            pass  # mixed types, keep original order
                        if paired:
                            actual_data_input = [p[0] for p in paired]
                            sorted_values = [p[1] for p in paired]
                            # Replace dict with sorted list for agent consumption
                            series_info["values"] = sorted_values
                            self.orch.current_metadata["series"] = series_info

                # === Generate unique analysis output directory ===
                # Deferred until after early-return checks (series variable
                # confirmation, missing series metadata) to avoid creating
                # orphan directories that never receive analysis results.
                analysis_id = self.orch.generate_analysis_id(data_path, agent_id)
                analysis_output_dir = self.orch.results_dir / f"analysis_{analysis_id}"
                analysis_output_dir.mkdir(parents=True, exist_ok=True)

                print(f"    Analysis ID: {analysis_id}")
                print(f"    Output directory: {analysis_output_dir}")

                # === Save metadata copy for traceability ===
                metadata_copy_path = analysis_output_dir / "metadata_used.json"
                with open(metadata_copy_path, 'w') as f:
                    json.dump({
                        "analysis_id": analysis_id,
                        "data_path": data_path,
                        "agent_id": agent_id,
                        "agent_name": self.AGENT_NAMES.get(agent_id),
                        "analysis_goal": analysis_goal,
                        "timestamp": datetime.now().isoformat(),
                        "metadata": _structure_metadata_for_save(
                            self.orch.current_metadata
                        ),
                    }, f, indent=2)

                # === Create agent with unique output directory ===
                # NOTE: Code-executing agents may prompt the user
                # for sandbox approval and raise RuntimeError if declined.
                try:
                    agent = self.orch.create_agent_for_analysis(agent_id, str(analysis_output_dir))
                except RuntimeError as e:
                    # Handle sandbox rejection or other init failures
                    error_msg = str(e)

                    if "sandbox" in error_msg.lower() or "declined" in error_msg.lower():
                        # Clean up the output directory we created
                        import shutil
                        if analysis_output_dir.exists():
                            shutil.rmtree(analysis_output_dir)

                        return json.dumps({
                            "status": "aborted",
                            "reason": "sandbox_declined",
                            "message": "Analysis aborted: User declined to proceed without sandbox protection.",
                            "agent_id": agent_id,
                            "agent_name": self.AGENT_NAMES.get(agent_id),
                            "recommendation": (
                                "This agent executes AI-generated Python code and requires a secure environment.\n\n"
                                "Please run SciLink in one of the following:\n"
                                "  1. Docker container (recommended)\n"
                                "  2. Virtual machine (VMware, VirtualBox, cloud VM)\n"
                                "  3. Google Colab\n\n"
                                "See the documentation for setup instructions."
                            )
                        })
                    else:
                        # Some other initialization error
                        raise

                print(f"    Using agent: {type(agent).__name__}")
                print(f"    Data: {data_path}")

                # === Run analysis ===
                analyze_kwargs = {
                    "data": actual_data_input,
                    "system_info": self.orch.current_metadata,
                }
                if objective is not None:
                    analyze_kwargs["objective"] = objective
                if hints is not None:
                    analyze_kwargs["hints"] = hints
                if auxiliary_data is not None:
                    analyze_kwargs["auxiliary_data"] = auxiliary_data
                if auxiliary_label is not None:
                    analyze_kwargs["auxiliary_label"] = auxiliary_label
                if skill is not None:
                    # PR 3: ``skill`` may be a single name or a list. Resolve
                    # any custom-skill names to their registered file paths
                    # so the agent's load_skill() can locate them.
                    custom_skills = getattr(self.orch, "_custom_skills", {})

                    def _resolve_one(s):
                        return custom_skills[s] if s in custom_skills else s

                    if isinstance(skill, str):
                        analyze_kwargs["skill"] = _resolve_one(skill)
                    elif isinstance(skill, (list, tuple)):
                        analyze_kwargs["skill"] = [_resolve_one(s) for s in skill]
                    else:
                        analyze_kwargs["skill"] = skill
                if task_mode is not None:
                    # Currently consumed by CurveFittingAgent; other agents
                    # accept **kwargs and silently ignore unknown parameters,
                    # matching the existing pattern for `hints`.
                    analyze_kwargs["task_mode"] = task_mode
                if self.orch.active_knowledge:
                    analyze_kwargs["prior_knowledge"] = self.orch.active_knowledge
                if prior_analysis_paths:
                    analyze_kwargs["prior_analysis_paths"] = prior_analysis_paths
                if literature_file:
                    analyze_kwargs["literature_file"] = literature_file
                result = agent.analyze(**analyze_kwargs)
                
                # === Store result ===
                analysis_record = {
                    "analysis_id": analysis_id,
                    "timestamp": datetime.now().isoformat(),
                    "data_path": data_path,
                    "agent_id": agent_id,
                    "agent_name": self.AGENT_NAMES.get(agent_id),
                    "status": result.get("status"),
                    "output_directory": str(analysis_output_dir),
                    "literature_file": literature_file,
                    "full_result": result,
                    "novelty_assessment": None
                }
                self.orch.analysis_results.append(analysis_record)
                
                # === Format response ===
                if result.get("status") == "success":
                    # Find main visualization
                    viz_path = None
                    for candidate in analysis_output_dir.rglob("*_analysis.png"):
                        viz_path = str(candidate)
                        break
                    if not viz_path:
                        for candidate in analysis_output_dir.rglob("*.png"):
                            if "report" not in candidate.name.lower():
                                viz_path = str(candidate)
                                break

                    response = {
                        "status": "success",
                        "analysis_id": analysis_id,
                        "agent_used": self.AGENT_NAMES.get(agent_id),
                        "output_directory": str(analysis_output_dir),
                        "detailed_analysis": result.get("detailed_analysis", "")[:2000],
                        "claims_count": len(result.get("scientific_claims", [])),
                        "full_result_available": True,
                        "note": f"All outputs saved to: {analysis_output_dir}",
                        "next_steps": "Use assess_novelty to check literature for these claims, or get_recommendations for follow-up experiments.",
                    }
                    if viz_path:
                        response["visualization_path"] = viz_path
                    if result.get("tier2_results"):
                        response["tier2_ran"] = True
                        t2 = result["tier2_results"]
                        response["tier2_focus"] = t2.get(
                            "analysis_approach", "deeper analysis"
                        )
                    # Emit a flat feature table (per-unit conditions + extracted
                    # scalar features) so downstream planning / BO can ingest the
                    # results as a file rather than re-typed prose.
                    feature_table = write_feature_table(analysis_output_dir)
                    if feature_table:
                        response["feature_table"] = feature_table
                    # #172: surface the locked-script reuse verdict so the
                    # orchestrator can act on a non-`good` outcome (a poorly
                    # fitting reused recipe, or a re-derived schema).
                    reuse_validity = result.get("reuse_validity")
                    if reuse_validity:
                        response["reuse_validity"] = reuse_validity
                        if reuse_validity.get("verdict") != "good":
                            response["reuse_warning"] = reuse_validity.get(
                                "message", ""
                            )
                    return json.dumps(response)
                else:
                    return json.dumps({
                        "status": "error",
                        "analysis_id": analysis_id,
                        "error": result.get("error", {}),
                        "agent_used": self.AGENT_NAMES.get(agent_id),
                        "output_directory": str(analysis_output_dir)
                    })
                    
            except Exception as e:
                self.logger.error(f"Analysis error: {e}", exc_info=True)
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=run_analysis,
            name="run_analysis",
            description=(
                "Execute analysis with the selected or specified agent. "
                "Each run creates a unique output directory (analysis_{dataset_name}_{timestamp}) "
                "for traceability. Requires data path and metadata to be set. "
                "For series analysis, the system resolves the control variable in order: "
                "(1) explicit series_metadata parameter, "
                "(2) automatic extraction from per-file JSON sidecars via LLM reasoning, "
                "(3) user prompt. "
                "Optional objective provides a high-level scientific question to frame the analysis "
                "(e.g. 'Determine the oxidation state of Ti'). "
                "Optional hints provide tactical guidance to steer the analysis "
                "(e.g. 'focus on the Ti L-edge around 460 eV'). "
                "Optional auxiliary_data provides a complementary dataset "
                "(e.g. TGA alongside DSC, or microscopy image) as context. "
                "Supported by CurveFitting and Hyperspectral agents. "
                "Optional skill provides domain-specific knowledge "
                "(e.g. 'xps', 'xrd') for improved fitting and interpretation. "
                "Returns analysis_id and output_directory for reference."
            ),
            parameters={
                "data_path": {
                    "type": "string",
                    "description": "Path to data file (uses current if not specified)"
                },
                "agent_id": {
                    "type": "integer",
                    "description": "Agent ID to use (0-3, uses selected if not specified)"
                },
                "analysis_goal": {
                    "type": "string",
                    "description": "Specific analysis objective (saved with results for traceability)"
                },
                "objective": {
                    "type": "string",
                    "description": (
                        "High-level scientific objective that frames the analysis "
                        "(e.g. 'Determine whether the sample underwent a phase transition', "
                        "'Quantify relative concentration of anatase vs rutile'). "
                        "Unlike hints, this tells the agent *why* the analysis is being "
                        "performed and *what question* to answer."
                    )
                },
                "hints": {
                    "type": "string",
                    "description": (
                        "Tactical guidance to steer the analysis "
                        "(e.g. 'focus on the Ti L-edge around 460 eV', "
                        "'pay attention to peaks between 280-300 nm'). "
                        "Supported by CurveFitting and Hyperspectral agents."
                    )
                },
                "auxiliary_data": {
                    "type": "string",
                    "description": (
                        "Path to an auxiliary dataset (1D curve or image) as visual "
                        "context for the analysis. Can also be a visualization from "
                        "a prior analysis. Only use this when the datasets are "
                        "genuinely related (same sample, same region, or "
                        "complementary techniques) — do not chain unrelated analyses."
                    )
                },
                "auxiliary_label": {
                    "type": "string",
                    "description": (
                        "Description of auxiliary data and its relationship to the "
                        "primary dataset, e.g. 'TGA curve collected simultaneously "
                        "during DSC' or 'SEM analysis of the same region'"
                    )
                },
                "skill": {
                    # Accepts a single skill name (string) or a list of names
                    # for multi-skill loading. Schema permits both shapes via
                    # JSON Schema ``oneOf``.
                    "oneOf": [
                        {"type": "string"},
                        {"type": "array", "items": {"type": "string"}},
                    ],
                    "description": _build_skill_description(
                        getattr(self.orch, "_agent_registry", None),
                        getattr(self.orch, "_custom_skills", None),
                    ),
                },
                "series_metadata": {
                    "type": "string",
                    "description": (
                        "JSON string describing the experimental variable that changes across "
                        "spectra in a series. Takes highest priority — overrides automatic "
                        "extraction from per-file JSON sidecars. "
                        "Values is a dict mapping each filename to its value — files are "
                        "automatically sorted by value for correct trend analysis. "
                        "Format: {\"variable\": \"<variable>\", \"values\": {\"<filename>\": <value>, ...}, \"unit\": \"<units>\"}. "
                        "Example: {\"variable\": \"temperature\", \"values\": {\"spec_5K.csv\": 5, \"spec_10K.csv\": 10, \"spec_20K.csv\": 20}, \"unit\": \"K\"}"
                    )
                },
                "task_mode": {
                    "type": "string",
                    "enum": ["fitting", "identification"],
                    "description": (
                        "CurveFitting agent only. Set to 'identification' when the user "
                        "is asking the agent to help identify what material or phase the "
                        "spectrum is from (no sample identity known), rather than to fit "
                        "and interpret a known material. In identification mode the planner "
                        "uses a generic flexible model and the interpreter enumerates "
                        "ranked candidate materials with discriminating peaks instead of "
                        "asserting a single answer. Leave unset (defaults to 'fitting') "
                        "for standard analyses where the sample is known."
                    )
                },
                "prior_analysis_paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": (
                        "List of folder or file paths from previous analyses whose "
                        "outputs the new run can consume. Use this when a follow-up "
                        "analysis needs to load artifacts (masks, positions, "
                        "feature tables, abundance maps) from a prior run rather "
                        "than recomputing them. Directory paths typically come "
                        "from `list_results()` (the `output_directory` field). "
                        "For each path, the agent's code generator receives the "
                        "file listing (loadable via absolute path); for paths "
                        "containing `analysis_results.json`, the planner also "
                        "receives a state summary (pipeline, quality score, "
                        "extracted features, scientific claims, saved-arrays "
                        "catalog). Consumed by the image-analysis agent and "
                        "the curve-fitting agent — for a prior curve-fit run "
                        "the saved fitting script and fit summary are surfaced "
                        "to its planning and script-generation stages.\n"
                        "LOCKED EXTRACTION-SCRIPT REUSE: when the new data is "
                        "the NEXT MEASUREMENT of the SAME measurement series as "
                        "the prior run — the same kind of unit, only the "
                        "control parameters differ (a new point in a "
                        "Bayesian-optimization / closed-loop campaign) — the "
                        "agent reuses the prior run's locked extraction script "
                        "verbatim instead of re-deriving the model/pipeline. "
                        "This guarantees the new feature row has the SAME "
                        "columns as the campaign, which the planning-side "
                        "feature-table append strictly requires. ONLY pass "
                        "prior_analysis_paths when the new data genuinely "
                        "continues that series; if it is a different kind of "
                        "measurement, do NOT pass it (a fresh analysis is "
                        "correct — forcing the prior recipe would be wrong). "
                        "The run reports a `reuse_validity` verdict "
                        "(`good` / `poor` / `script_failed`) — read it and act "
                        "on a non-`good` verdict (see the system guidance)."
                    )
                },
                "literature_file": {
                    "type": "string",
                    "description": (
                        "Path to a markdown file of literature / document "
                        "context — produced by `search_literature` (external "
                        "search) or `read_document` (user-provided papers). "
                        "When provided, its contents are injected into the "
                        "planner so the proposed analysis plan is grounded in "
                        "that literature. Skipped for the curve-fitting agent's "
                        "`task_mode='identification'` to preserve the unbiased "
                        "fit; in that case the literature still informs Stage-2 "
                        "candidate enumeration."
                    )
                }
            },
            required=[]
        )

        # =====================================================================
        # 6. LIST RESULTS
        # =====================================================================
        def list_results() -> str:
            """
            List analysis results in the session directory.
            Shows all analysis runs with their IDs and output directories.
            """
            print(f"  ⚡ Tool: Listing results...")
            
            results = []
            
            # List analysis directories in results folder
            results_dir = self.orch.results_dir
            if results_dir.exists():
                # Find all analysis directories
                analysis_dirs = sorted(
                    [d for d in results_dir.iterdir() if d.is_dir() and d.name.startswith("analysis_")],
                    key=lambda x: x.stat().st_mtime,
                    reverse=True  # Most recent first
                )
                
                for analysis_dir in analysis_dirs:
                    analysis_info = {
                        "directory": analysis_dir.name,
                        "path": str(analysis_dir),
                        "files": [],
                        "has_novelty_assessment": False
                    }
                    
                    # Check for metadata_used.json to get analysis details
                    metadata_file = analysis_dir / "metadata_used.json"
                    if metadata_file.exists():
                        try:
                            with open(metadata_file, 'r') as f:
                                meta = json.load(f)
                            analysis_info["analysis_id"] = meta.get("analysis_id")
                            analysis_info["data_path"] = meta.get("data_path")
                            analysis_info["agent_name"] = meta.get("agent_name")
                            analysis_info["timestamp"] = meta.get("timestamp")
                        except Exception:
                            pass
                    
                    # Check for novelty assessment
                    novelty_file = analysis_dir / "literature_assessment" / "novelty_report.json"
                    if novelty_file.exists():
                        analysis_info["has_novelty_assessment"] = True
                    
                    # List files in directory
                    for f in analysis_dir.iterdir():
                        if f.is_file():
                            analysis_info["files"].append(f.name)
                    
                    results.append(analysis_info)
            
            # Also include in-memory analysis history
            return json.dumps({
                "status": "success",
                "session_directory": str(self.orch.base_dir),
                "results_directory": str(results_dir),
                "total_analyses": len(results),
                "analyses": results,
                "in_memory_history": [
                    {
                        "analysis_id": r.get("analysis_id"),
                        "data_path": r.get("data_path"),
                        "agent_name": r.get("agent_name"),
                        "status": r.get("status"),
                        "output_directory": r.get("output_directory"),
                        "has_novelty_assessment": r.get("novelty_assessment") is not None
                    }
                    for r in self.orch.analysis_results
                ]
            })
        
        self._register_tool(
            func=list_results,
            name="list_results",
            description=(
                "List all analysis results in the session. "
                "Shows analysis IDs, data paths, agents used, and output directories."
            ),
            parameters={},
            required=[]
        )
        
        # =====================================================================
        # 7. SAVE CHECKPOINT
        # =====================================================================
        def save_checkpoint() -> str:
            """
            Save session state for later resumption.
            """
            print(f"  ⚡ Tool: Saving checkpoint...")
            
            try:
                checkpoint_data = {
                    "timestamp": datetime.now().isoformat(),
                    "current_metadata": self.orch.current_metadata,
                    "current_data_path": self.orch.current_data_path,
                    "current_data_type": self.orch.current_data_type,
                    "selected_agent_id": self.orch.selected_agent_id,
                    "analysis_results": self.orch.analysis_results,
                    "analysis_run_counter": self.orch._analysis_run_counter,
                    "message_count": self.orch.message_count,
                    "analysis_mode": self.orch.analysis_mode.value,
                    "active_knowledge": self.orch.active_knowledge,
                    "graduated_skill_sources": self.orch._graduated_skill_sources,
                }

                with open(self.orch.checkpoint_path, 'w') as f:
                    json.dump(checkpoint_data, f, indent=2)

                return json.dumps({
                    "status": "success",
                    "checkpoint_path": str(self.orch.checkpoint_path),
                    "timestamp": checkpoint_data["timestamp"],
                    "analyses_saved": len(self.orch.analysis_results)
                })
                
            except Exception as e:
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=save_checkpoint,
            name="save_checkpoint",
            description=(
                "Save complete session state including metadata, analysis history, "
                "and conversation. Use this to enable session resumption."
            ),
            parameters={},
            required=[]
        )
        
        # =====================================================================
        # 8. SHOW AVAILABLE AGENTS
        # =====================================================================
        def show_available_agents() -> str:
            """
            Show list of available analysis agents and their capabilities,
            plus any custom tools registered via register_tools().
            """
            print(f"  ⚡ Tool: Showing available agents...")

            agents = []
            for agent_id in sorted(self.AGENT_NAMES.keys()):
                agents.append({
                    "id": agent_id,
                    "name": self.AGENT_NAMES[agent_id],
                    "description": self.AGENT_DESCRIPTIONS[agent_id]
                })

            available_skills = list_all_skills()
            custom_skills = getattr(self.orch, "_custom_skills", {})
            if custom_skills:
                available_skills["custom"] = sorted(custom_skills.keys())

            result = {
                "status": "success",
                "agents": agents,
                "current_selection": self.orch.selected_agent_id,
                "available_skills": available_skills,
            }

            external_tools = getattr(self.orch, "_external_tools", [])
            if external_tools:
                result["custom_tools"] = external_tools
                result["custom_tools_note"] = (
                    "These tools are callable directly by name and operate on "
                    "the current data file (set via examine_data)."
                )

            return json.dumps(result)

        self._register_tool(
            func=show_available_agents,
            name="show_available_agents",
            description=(
                "Show list of available analysis agents and their capabilities, "
                "plus any custom tools registered for the current session."
            ),
            parameters={},
            required=[]
        )
        
        # =====================================================================
        # 9. GET METADATA SCHEMA
        # =====================================================================
        def get_metadata_schema() -> str:
            """
            Get the metadata JSON schema for reference.
            """
            print(f"  ⚡ Tool: Getting metadata schema...")
            
            return json.dumps({
                "status": "success",
                "schema": METADATA_SCHEMA_DICT,
                "required_fields": ["experiment_type", "experiment", "sample"],
                "hint": "Use convert_metadata to create metadata from natural language"
            })
        
        self._register_tool(
            func=get_metadata_schema,
            name="get_metadata_schema",
            description=(
                "Get the metadata JSON schema showing required and optional fields. "
                "Use this to understand what metadata is needed."
            ),
            parameters={},
            required=[]
        )
        
        # =====================================================================
        # 10. GET MEASUREMENT RECOMMENDATIONS (UPDATED)
        # =====================================================================
        def get_recommendations(analysis_id: str = None, analysis_index: int = -1) -> str:
            """
            Get measurement recommendations from a completed analysis.
            Can specify by analysis_id or by index in the history.
            
            UPDATED: Now incorporates novelty assessment results to prioritize
            recommendations based on scientific novelty.
            """
            print(f"  ⚡ Tool: Getting measurement recommendations...")
            
            if not self.orch.analysis_results:
                return json.dumps({
                    "status": "error",
                    "message": "No analyses completed yet. Run an analysis first."
                })
            
            try:
                # Find the analysis record
                record = None
                
                if analysis_id:
                    # Search by analysis_id
                    for r in self.orch.analysis_results:
                        if r.get("analysis_id") == analysis_id:
                            record = r
                            break
                    if record is None:
                        return json.dumps({
                            "status": "error",
                            "message": f"Analysis not found: {analysis_id}"
                        })
                else:
                    # Use index
                    record = self.orch.analysis_results[analysis_index]
                
                agent_id = record.get("agent_id")
                if agent_id is None:
                    return json.dumps({
                        "status": "error",
                        "message": "Analysis record missing agent_id"
                    })
                
                # Get the stored analysis result
                full_result = record.get("full_result")
                if full_result is None:
                    return json.dumps({
                        "status": "error",
                        "message": "Analysis result not stored. Please run the analysis again."
                    })
                
                # Get novelty assessment if available
                novelty_assessment = record.get("novelty_assessment")
                
                # Create agent for recommendations (uses same output dir)
                output_dir = record.get("output_directory", str(self.orch.results_dir / "temp"))
                agent = self.orch.create_agent_for_analysis(agent_id, output_dir)
                
                # Call recommend_measurements with the stored result AND novelty assessment
                result = agent.recommend_measurements(
                    data=record.get("data_path"),
                    system_info=self.orch.current_metadata,
                    analysis_result=full_result,
                    novelty_assessment=novelty_assessment  # NEW: Pass novelty data
                )
                
                response = {
                    "status": result.get("status", "success"),
                    "analysis_id": record.get("analysis_id"),
                    "recommendations": result.get("measurement_recommendations", []),
                    "analysis_integration": result.get("analysis_integration", ""),
                    "novelty_informed": novelty_assessment is not None
                }
                
                # Add novelty-specific recommendations if available
                if novelty_assessment:
                    response["novelty_summary"] = {
                        "total_claims_assessed": len(novelty_assessment.get("assessments", [])),
                        "high_novelty_claims": len(novelty_assessment.get("high_novelty_claims", [])),
                        "novelty_driven_recommendations": result.get("novelty_recommendations", [])
                    }
                
                return json.dumps(response)
                
            except Exception as e:
                self.logger.error(f"Recommendations error: {e}", exc_info=True)
                return json.dumps({
                    "status": "error",
                    "message": str(e)
                })
        
        self._register_tool(
            func=get_recommendations,
            name="get_recommendations",
            description=(
                "Get measurement recommendations based on a completed analysis. "
                "Specify by analysis_id or use analysis_index (-1 for most recent). "
                "Returns suggested follow-up experiments and measurements. "
                "If assess_novelty was run first, recommendations are prioritized "
                "based on scientific novelty (high-novelty claims get validation experiments)."
            ),
            parameters={
                "analysis_id": {
                    "type": "string",
                    "description": "Specific analysis ID to get recommendations for"
                },
                "analysis_index": {
                    "type": "integer",
                    "description": "Index of analysis in history (-1 for most recent)"
                }
            },
            required=[]
        )

        # =====================================================================
        # 10b. SEARCH LITERATURE (preparatory — call BEFORE run_analysis)
        # =====================================================================
        def search_literature(query: str) -> str:
            """
            Search scientific literature for context to inform an upcoming
            analysis. Returns a file path that should be passed as
            `literature_file` to the next `run_analysis` call so the planner
            can produce a literature-informed plan.
            """
            print(f"  ⚡ Tool: Searching literature for '{query[:80]}...'")

            if not self.orch.futurehouse_api_key:
                return json.dumps({
                    "status": "error",
                    "message": "No FutureHouse/Edison API Key provided in Orchestrator initialization."
                })

            try:
                # 50-min ceiling matches plan mode's `LiteratureSearchAgent`
                # default; Edison CROW jobs can take 20-30 min for harder queries.
                lit_agent = FittingModelLiteratureAgent(
                    api_key=self.orch.futurehouse_api_key, max_wait_time=3000
                )
            except Exception as e:
                return json.dumps({"status": "error", "message": f"Failed to init Literature Agent: {e}"})

            # Refine the orchestrator-LLM's draft query using the loaded data
            # preview + metadata (visual + experimental specifics). Best-effort:
            # falls back to the raw query on any failure.
            refined_query = optimize_query_for_analysis(
                raw_query=query,
                data_type=getattr(self.orch, "current_data_type", None),
                data_path=getattr(self.orch, "current_data_path", None),
                metadata=getattr(self.orch, "current_metadata", None),
                model=self.orch.model,
            )
            if refined_query != query:
                print(f"  🔍 Refined query: {refined_query}")
                logging.info(f"search_literature: refined query → {refined_query}")
            else:
                print(f"  🔍 Using raw query (no refinement applied)")

            try:
                result = lit_agent.query_for_models(refined_query)
            except Exception as e:
                logging.error(f"Literature search error: {e}", exc_info=True)
                return json.dumps({"status": "error", "message": str(e)})

            if result.get("status") != "success":
                return json.dumps({
                    "status": result.get("status", "error"),
                    "message": result.get("message", "Literature search did not succeed")
                })

            content = result.get("formatted_answer", "") or ""

            # Hash the raw query for an idempotent, collision-free filename:
            # same raw query → same file (re-runnable), different queries → different files.
            query_hash = hashlib.md5(query.encode("utf-8")).hexdigest()[:8]
            lit_path = self.orch.base_dir / f"literature_search_{query_hash}.md"
            with open(lit_path, "w") as f:
                f.write(f"# Literature Search Results\n\n**Draft query:** {query}\n")
                if refined_query != query:
                    f.write(f"**Refined query:** {refined_query}\n")
                f.write(f"\n{content}")

            print(f"  ✅ Literature search completed. Saved to {lit_path.name}")

            preview = content[:500] + "..." if len(content) > 500 else content
            return json.dumps({
                "status": "success",
                "file_path": str(lit_path),
                "content_preview": preview,
                "hint": "Pass file_path as literature_file to the next run_analysis() call."
            })

        self._register_tool(
            func=search_literature,
            name="search_literature",
            description=(
                "Search scientific literature via the FutureHouse Edison API to gather "
                "context that will inform an upcoming analysis. Call BEFORE run_analysis(); "
                "pass the returned file_path as `literature_file` to run_analysis() so the "
                "planner produces a literature-informed plan."
            ),
            parameters={
                "query": {
                    "type": "string",
                    "description": "A focused research question (e.g., 'methods for detecting grain boundaries in HRTEM images of 2D materials')."
                }
            },
            required=["query"]
        )

        # =====================================================================
        # 11. ASSESS NOVELTY
        # =====================================================================
        def assess_novelty(analysis_id: str = None, analysis_index: int = -1) -> str:
            """
            Perform a literature search and novelty assessment on claims generated 
            by a previous analysis run.
            
            UPDATED: Now stores results in the analysis record for use by
            get_recommendations.
            """
            print(f"  ⚡ Tool: Assessing novelty for analysis...")

            if not self.orch.futurehouse_api_key:
                return json.dumps({
                    "status": "error",
                    "message": "No FutureHouse/Edison API Key provided in Orchestrator initialization."
                })

            # 1. Retrieve the analysis record
            record = None
            record_index = None
            
            if analysis_id:
                for i, r in enumerate(self.orch.analysis_results):
                    if r.get("analysis_id") == analysis_id:
                        record = r
                        record_index = i
                        break
                if record is None:
                    return json.dumps({"status": "error", "message": f"Analysis ID not found: {analysis_id}"})
            else:
                if not self.orch.analysis_results:
                    return json.dumps({"status": "error", "message": "No analysis history available."})
                record_index = analysis_index if analysis_index >= 0 else len(self.orch.analysis_results) + analysis_index
                record = self.orch.analysis_results[record_index]

            # 2. Extract Claims
            full_result = record.get("full_result", {})
            claims = full_result.get("scientific_claims", [])
            
            if not claims:
                return json.dumps({
                    "status": "warning",
                    "message": "No scientific claims found in this analysis to assess."
                })

            print(f"    Found {len(claims)} claims to assess from {record.get('analysis_id')}")

            # 3. Initialize Lit Agents
            try:
                owl_agent = OwlLiteratureAgent(api_key=self.orch.futurehouse_api_key, max_wait_time=600)
                
                # Use orchestrator's generic LLM config for the Scorer
                scorer = NoveltyScorer(
                    api_key=self.orch.api_key,
                    model_name=self.orch.model_name,
                    base_url=self.orch.base_url
                )
            except Exception as e:
                return json.dumps({"status": "error", "message": f"Failed to init Lit Agents: {e}"})

            # 4. Process Claims
            scored_results = []
            high_novelty_claims = []
            
            # Create a dedicated directory for lit results inside the analysis folder
            analysis_dir = Path(record.get("output_directory", self.orch.results_dir))
            lit_output_dir = analysis_dir / "literature_assessment"
            lit_output_dir.mkdir(exist_ok=True)

            print(f"    Output directory: {lit_output_dir}")

            for i, claim_obj in enumerate(claims):
                question = claim_obj.get("has_anyone_question")
                claim_text = claim_obj.get("claim")
                
                if not question:
                    continue

                print(f"    🔍 Searching claim {i+1}/{len(claims)}: {question[:60]}...")
                
                # Search (Owl)
                search_res = owl_agent.query_literature(question)
                
                if search_res.get("status") != "success":
                    print(f"       ⚠️ Search failed for claim {i+1}")
                    continue

                formatted_answer = search_res.get("formatted_answer", "")

                # Score (Scorer)
                print(f"       ⚖️ Scoring novelty...")
                score_res = scorer.score_novelty(question, formatted_answer)
                
                novelty_score = score_res.get("novelty_score", 0)
                
                result_entry = {
                    "claim_index": i,
                    "original_claim": claim_text,
                    "question": question,
                    "search_answer": formatted_answer,
                    "novelty_score": novelty_score,
                    "novelty_explanation": score_res.get("explanation"),
                    "sources": [s.url for s in getattr(search_res, 'sources', []) if hasattr(s, 'url')]
                }
                scored_results.append(result_entry)
                
                # Track high-novelty claims for recommendations
                if novelty_score >= 4:
                    high_novelty_claims.append(result_entry)
                
                # Pause briefly to be polite to APIs
                time.sleep(1)

            # 5. Build novelty assessment object
            novelty_assessment = {
                "timestamp": datetime.now().isoformat(),
                "assessments": scored_results,
                "high_novelty_claims": high_novelty_claims,
                "summary_stats": {
                    "total_assessed": len(scored_results),
                    "high_novelty_count": len(high_novelty_claims),
                    "average_score": sum(r.get("novelty_score", 0) for r in scored_results) / len(scored_results) if scored_results else 0
                }
            }
            
            # 6. Store in the analysis record (KEY CHANGE)
            self.orch.analysis_results[record_index]["novelty_assessment"] = novelty_assessment
            
            # 7. Save Results to file
            output_file = lit_output_dir / "novelty_report.json"
            with open(output_file, "w") as f:
                json.dump({
                    "analysis_id": record.get("analysis_id"),
                    **novelty_assessment
                }, f, indent=2)

            # 8. Summarize for Chat
            summary_lines = []
            
            for res in scored_results:
                score = res['novelty_score']
                if score >= 4:
                    icon = "🌟"
                elif score == 3:
                    icon = "🤔"
                else:
                    icon = "📚"
                
                summary_lines.append(
                    f"{icon} [Score {score}/5] {res['original_claim'][:50]}... "
                    f"-> {res['novelty_explanation'][:80]}..."
                )

            return json.dumps({
                "status": "success",
                "total_assessed": len(scored_results),
                "high_novelty_count": len(high_novelty_claims),
                "average_novelty_score": novelty_assessment["summary_stats"]["average_score"],
                "summary_text": "\n".join(summary_lines),
                "report_path": str(output_file),
                "stored_for_recommendations": True,
                "note": "Novelty assessment stored. Use get_recommendations to get novelty-informed follow-up suggestions."
            })

        self._register_tool(
            func=assess_novelty,
            name="assess_novelty",
            description=(
                "Perform a literature search to assess the novelty of scientific claims "
                "generated by a previous analysis. Requires an analysis_id (from run_analysis). "
                "Returns novelty scores (1-5) and checks for prior art. "
                "Results are stored and used by get_recommendations for prioritized suggestions."
            ),
            parameters={
                "analysis_id": {
                    "type": "string",
                    "description": "The ID of the analysis run to assess (e.g. 'sample1_FFT_2023...')"
                },
                "analysis_index": {
                    "type": "integer",
                    "description": "Alternatively, use the index of the analysis in memory (-1 for most recent)"
                }
            },
            required=[]
        )

        # =====================================================================
        # 11b. RECOMMEND DFT STRUCTURES
        # =====================================================================
        def recommend_dft_structures(analysis_id: str = None,
                                     analysis_index: int = -1) -> str:
            """
            Generate DFT structure recommendations from a completed analysis,
            optionally informed by a prior novelty assessment.
            """
            print(f"  ⚡ Tool: Generating DFT structure recommendations...")

            # 1. Locate the analysis record
            record = None
            record_index = None

            if analysis_id:
                for i, r in enumerate(self.orch.analysis_results):
                    if r.get("analysis_id") == analysis_id:
                        record = r
                        record_index = i
                        break
                if record is None:
                    return json.dumps({"status": "error", "message": f"Analysis ID not found: {analysis_id}"})
            else:
                if not self.orch.analysis_results:
                    return json.dumps({"status": "error", "message": "No analysis history available."})
                record_index = analysis_index if analysis_index >= 0 else len(self.orch.analysis_results) + analysis_index
                record = self.orch.analysis_results[record_index]

            # 2. Extract analysis text
            full_result = record.get("full_result") or {}
            analysis_text = (
                full_result.get("detailed_analysis")
                or full_result.get("full_analysis")
                or ""
            )
            if not analysis_text:
                return json.dumps({
                    "status": "error",
                    "message": "Analysis record has no detailed_analysis text to work from."
                })

            # 3. Pull novel-claim *strings* (high_novelty_claims entries are dicts)
            novelty = record.get("novelty_assessment") or {}
            novel_claim_dicts = novelty.get("high_novelty_claims", []) or []
            novel_claims = [
                c.get("original_claim", "") if isinstance(c, dict) else str(c)
                for c in novel_claim_dicts
            ]
            novel_claims = [c for c in novel_claims if c]

            # 4. Build novelty context (mirror DFTRecommender.run_from_data)
            if novel_claims:
                context = "Focus on these potentially novel findings:\n"
                for i, claim in enumerate(novel_claims, 1):
                    context += f"{i}. {claim}\n"
                context += "\nPrioritize DFT structures that can investigate these novel aspects."
            else:
                context = "No specific novel claims identified. Focus on most interesting aspects."

            # 5. Output dir nested under the analysis record's directory
            base_dir = Path(record.get("output_directory", self.orch.results_dir))
            out_dir = base_dir / "dft_recommendations"
            out_dir.mkdir(parents=True, exist_ok=True)

            # 6. Run RecommendationAgent directly (no need for the DFTRecommender wrapper)
            try:
                agent = RecommendationAgent(
                    api_key=self.orch.api_key,
                    base_url=self.orch.base_url,
                    model_name=self.orch.model_name,
                )
                result = agent.generate_dft_recommendations_from_text(
                    cached_detailed_analysis=analysis_text,
                    additional_prompt_context=context,
                    system_info=self.orch.current_metadata,
                )
            except Exception as e:
                return json.dumps({"status": "error", "message": f"Failed to generate recommendations: {e}"})

            if "error" in result:
                return json.dumps({"status": "error", "message": result.get("error")})

            recommendations = result.get("recommendations", []) or []
            reasoning = result.get("analysis_summary_or_reasoning", "")

            # 7. Persist sidecar JSON for parity with the standalone runner
            output_file = out_dir / "dft_recommendations.json"
            try:
                with open(output_file, 'w') as f:
                    json.dump({
                        "reasoning": reasoning,
                        "recommendations": recommendations,
                        "novel_claims": novel_claims,
                    }, f, indent=2)
            except Exception as e:
                self.logger.warning(f"Failed to write DFT recommendations sidecar: {e}")

            # 8. Persist on the analysis record for downstream tools
            self.orch.analysis_results[record_index]["dft_recommendations"] = recommendations

            print(f"    Generated {len(recommendations)} DFT recommendations → {output_file}")

            return json.dumps({
                "status": "success",
                "count": len(recommendations),
                "recommendations": [
                    {
                        "priority": r.get("priority"),
                        "description": r.get("description"),
                        "scientific_interest": r.get("scientific_interest"),
                    }
                    for r in recommendations
                ],
                "output_file": str(output_file),
            })

        self._register_tool(
            func=recommend_dft_structures,
            name="recommend_dft_structures",
            description=(
                "Generate DFT structure recommendations from a completed analysis. "
                "Specify by analysis_id or use analysis_index (-1 for most recent). "
                "If assess_novelty was run first, recommendations focus on novel claims. "
                "Stores recommendations on the analysis record for use by run_dft_workflow."
            ),
            parameters={
                "analysis_id": {
                    "type": "string",
                    "description": "The ID of the analysis run to use (e.g. 'sample1_FFT_2023...')"
                },
                "analysis_index": {
                    "type": "integer",
                    "description": "Alternatively, use the index of the analysis in memory (-1 for most recent)"
                }
            },
            required=[]
        )

        # =====================================================================
        # 11c. RUN DFT WORKFLOW (DFTOrchestrator)
        # =====================================================================
        def run_dft_workflow(structure_description: str = None,
                             analysis_id: str = None,
                             analysis_index: int = -1,
                             recommendation_index: int = None,
                             vasp_generator_method: str = "atomate2",
                             max_refinement_cycles: int = 4) -> str:
            """
            Run the DFT orchestrator to produce VASP-ready inputs (POSCAR, INCAR,
            KPOINTS) for a given structure description, or for a structure picked
            from the recommendations stored on a previous analysis record.
            """
            print(f"  ⚡ Tool: Running DFT workflow...")

            # 1. Locate the optional analysis record
            record = None
            record_index = None
            if analysis_id:
                for i, r in enumerate(self.orch.analysis_results):
                    if r.get("analysis_id") == analysis_id:
                        record = r
                        record_index = i
                        break
                if record is None:
                    return json.dumps({"status": "error", "message": f"Analysis ID not found: {analysis_id}"})
            elif self.orch.analysis_results:
                record_index = analysis_index if analysis_index >= 0 else len(self.orch.analysis_results) + analysis_index
                if 0 <= record_index < len(self.orch.analysis_results):
                    record = self.orch.analysis_results[record_index]

            # 2. Resolve the structure description
            if recommendation_index is not None:
                if record is None:
                    return json.dumps({"status": "error",
                                       "message": "recommendation_index requires an analysis_id or available analysis history."})
                recs = record.get("dft_recommendations") or []
                if not (0 <= recommendation_index < len(recs)):
                    return json.dumps({"status": "error",
                                       "message": f"recommendation_index {recommendation_index} out of range (have {len(recs)})."})
                structure_description = recs[recommendation_index].get("description") or structure_description
            if not structure_description:
                return json.dumps({"status": "error",
                                   "message": "structure_description is required (or pass recommendation_index with an analysis that has stored recommendations)."})

            # 3. Build output directory
            slug = re.sub(r'[^A-Za-z0-9_-]+', '_', structure_description)[:40].strip('_') or "structure"
            base_dir = Path(record.get("output_directory")) if record else Path(self.orch.results_dir)
            out_dir = base_dir / "dft" / slug
            out_dir.mkdir(parents=True, exist_ok=True)

            # 4. Run the orchestrator (lazy import — keeps [sim] extras optional)
            try:
                from ..sim_agents.dft_orchestrator import DFTOrchestrator
            except ImportError as e:
                # vasp_generator_method='llm' only needs ase; 'atomate2' also
                # needs pymatgen + atomate2. Both paths require ase for
                # structure validation.
                if vasp_generator_method == "atomate2":
                    hint = "pip install ase pymatgen atomate2  (or: pip install 'scilink[sim]')"
                else:
                    hint = "pip install ase  (sufficient for vasp_generator_method='llm')"
                return json.dumps({
                    "status": "error",
                    "message": (
                        f"DFT workflow could not load required dependency. "
                        f"Install: {hint}. Original error: {e}"
                    ),
                })
            try:
                wf = DFTOrchestrator(
                    api_key=self.orch.api_key,
                    base_url=self.orch.base_url,
                    futurehouse_api_key=self.orch.futurehouse_api_key,
                    mp_api_key=None,  # DFTOrchestrator auto-discovers via get_api_key
                    generator_model=self.orch.model_name,
                    validator_model=self.orch.model_name,
                    output_dir=str(out_dir),
                    max_refinement_cycles=max_refinement_cycles,
                    vasp_generator_method=vasp_generator_method,
                )
                result = wf.run_complete_workflow(structure_description)
            except Exception as e:
                return json.dumps({"status": "error", "message": f"DFT workflow failed: {e}"})

            final_status = result.get("final_status")

            # 5. Persist on the analysis record
            if record is not None:
                self.orch.analysis_results[record_index].setdefault("dft_runs", []).append({
                    "description": structure_description,
                    "output_directory": str(out_dir),
                    "final_status": final_status,
                })

            # Surface refinement diagnostics to the agent. Without these the
            # agent sees only `ready_for_vasp: true` and reports success even
            # when the structure was accepted with substantial unresolved
            # validation issues (e.g., circuit-breaker fired on diverging
            # validator complaints).
            structure_gen = result.get("structure_generation", {}) or {}
            structure_warning = structure_gen.get("warning")
            cycles_used = structure_gen.get("cycles_used")
            val_result = structure_gen.get("validation_result", {}) or {}
            outstanding_issues = val_result.get("all_identified_issues", []) or []

            return json.dumps({
                "status": final_status if final_status else "error",
                "final_status": final_status,
                "output_directory": str(out_dir),
                "manifest_path": str(out_dir / "final_files_manifest.json"),
                "ready_for_vasp": final_status == "success",
                "structure_warning": structure_warning,
                "structure_refinement_cycles": cycles_used,
                "structure_outstanding_issues_count": len(outstanding_issues),
                "structure_outstanding_issues": outstanding_issues[:10],
            })

        self._register_tool(
            func=run_dft_workflow,
            name="run_dft_workflow",
            description=(
                "Run the DFT orchestrator to produce VASP-ready inputs (POSCAR, INCAR, "
                "KPOINTS) for a structure. Provide either an explicit structure_description, "
                "or an analysis_id + recommendation_index to pick a structure from prior "
                "recommend_dft_structures output. Does not run VASP itself; only generates inputs."
            ),
            parameters={
                "structure_description": {
                    "type": "string",
                    "description": "Free-text description of the structure to build (e.g. 'MoS2 monolayer with sulfur vacancy')."
                },
                "analysis_id": {
                    "type": "string",
                    "description": "Analysis run to attach this DFT job to (used for output dir + recommendation lookup)."
                },
                "analysis_index": {
                    "type": "integer",
                    "description": "Alternatively, index of the analysis in memory (-1 for most recent)."
                },
                "recommendation_index": {
                    "type": "integer",
                    "description": "Index into the analysis record's stored DFT recommendations to use as the structure description."
                },
                "vasp_generator_method": {
                    "type": "string",
                    "enum": ["llm", "atomate2"],
                    "description": "How to produce INCAR/KPOINTS. 'atomate2' is rule-based and fast; 'llm' is more flexible but slower."
                },
                "max_refinement_cycles": {
                    "type": "integer",
                    "description": "Maximum validator-guided structure refinement cycles."
                }
            },
            required=[]
        )

        # =====================================================================
        # 12. SET CUSTOM PREPROCESSING INSTRUCTION
        # =====================================================================
        def set_preprocessing_instruction(instruction: str, mode: str = "auto") -> str:
            """
            Add or update a custom preprocessing instruction in the current metadata.
            Metadata must already be loaded via load_metadata or convert_metadata.
            
            Modes:
                - "auto": If existing instruction found, return conflict for LLM to resolve
                - "replace": Overwrite existing instruction
                - "append": Append new instruction to existing one (blocked if redundant)
                - "force_append": Append without redundancy check
            """
            print(f"  ⚡ Tool: Setting custom preprocessing instruction...")
            
            if self.orch.current_metadata is None:
                return json.dumps({
                    "status": "error",
                    "message": "No metadata loaded. Use load_metadata or convert_metadata first."
                })

            # Treat empty/whitespace-only instruction as clearing custom preprocessing
            if not instruction or not instruction.strip():
                self.orch.current_metadata.pop("custom_processing_instruction", None)
                return json.dumps({
                    "status": "success",
                    "message": "Custom preprocessing instruction cleared (empty instruction)."
                })

            existing = self.orch.current_metadata.get("custom_processing_instruction")

            # Conflict detection
            if existing and mode == "auto":
                result = {
                    "status": "conflict",
                    "message": "Metadata already contains a custom preprocessing instruction.",
                    "existing_instruction": existing,
                    "new_instruction": instruction,
                    "options": [
                        "Call again with mode='replace' to overwrite the existing instruction.",
                        "Call again with mode='append' to combine both instructions.",
                    ],
                    "hint": "Ask the user which they prefer if unclear."
                }
                if self._check_instruction_redundancy(existing, instruction):
                    result["redundancy_warning"] = (
                        "These instructions appear to describe the same processing "
                        "operation. Appending would likely cause the same processing "
                        "to be applied twice, corrupting the data. Prefer 'replace' "
                        "unless you are certain they describe distinct steps."
                    )
                return json.dumps(result)

            if existing and mode == "append":
                if self._check_instruction_redundancy(existing, instruction):
                    return json.dumps({
                        "status": "conflict",
                        "message": (
                            "The new instruction appears to describe the same "
                            "processing as the existing one. Appending would likely "
                            "apply the same operation twice, corrupting the data."
                        ),
                        "existing_instruction": existing,
                        "new_instruction": instruction,
                        "options": [
                            "Call again with mode='replace' to use only the new instruction.",
                            "If you are certain these are distinct steps, call "
                            "again with mode='force_append' to combine them."
                        ],
                    })
                combined = f"{existing}\nThen: {instruction}"
                self.orch.current_metadata["custom_processing_instruction"] = combined
                return json.dumps({
                    "status": "success",
                    "message": "Appended new instruction to existing one.",
                    "final_instruction": combined
                })

            if existing and mode == "force_append":
                combined = f"{existing}\nThen: {instruction}"
                self.orch.current_metadata["custom_processing_instruction"] = combined
                return json.dumps({
                    "status": "success",
                    "message": "Force-appended new instruction to existing one.",
                    "final_instruction": combined
                })
            
            # mode == "replace" or no existing instruction
            self.orch.current_metadata["custom_processing_instruction"] = instruction
            
            result = {
                "status": "success",
                "message": "Custom preprocessing instruction set.",
                "instruction": instruction
            }
            if existing:
                result["note"] = f"Replaced previous instruction: '{existing}'"
            
            return json.dumps(result)

        self._register_tool(
            func=set_preprocessing_instruction,
            name="set_preprocessing_instruction",
            description=(
                "Add or update a custom DATA PREPROCESSING instruction in the currently loaded metadata. "
                "Use ONLY for raw data transformations BEFORE fitting: baseline division/subtraction, "
                "background correction, normalization, dark reference subtraction, smoothing, etc. "
                "Do NOT use for fitting model choices (e.g., 'use Lorentzian', 'fit with Fano', "
                "'fit the peak with a Voigt') — those go in the `hints` parameter of `run_analysis`. "
                "If metadata already has a preprocessing instruction, returns a conflict "
                "for you to resolve with the user. When appending, an LLM check detects "
                "redundant instructions to prevent double-processing. "
                "Supports modes: 'auto' (detect conflict), 'replace' (overwrite), "
                "'append' (combine both, with redundancy check), "
                "'force_append' (combine without redundancy check)."
            ),
            parameters={
                "instruction": {
                    "type": "string",
                    "description": (
                        "Natural language preprocessing instruction. Include file paths if "
                        "referencing external data."
                    )
                },
                "mode": {
                    "type": "string",
                    "description": (
                        "How to handle existing instructions: "
                        "'auto' (default, detect conflicts), "
                        "'replace' (overwrite), "
                        "'append' (combine both, blocks if redundant), "
                        "'force_append' (combine without redundancy check)"
                    )
                }
            },
            required=["instruction"]
        )

        # =====================================================================
        # 13. SYNTHESIZE KNOWLEDGE
        # =====================================================================
        def synthesize_knowledge(analysis_ids: list, focus: str, synthesis_type: str = "reference") -> str:
            """
            Distill findings from completed analyses into reusable knowledge.
            The synthesized knowledge is automatically injected into all
            subsequent run_analysis calls.
            """
            from scilink.knowledge import synthesize_knowledge as _synthesize

            print(f"  ⚡ Tool: Synthesizing knowledge ({synthesis_type}) from {len(analysis_ids)} analyses...")

            # Collect result dicts by analysis ID
            results = []
            missing_ids = []
            for aid in analysis_ids:
                found = False
                for record in self.orch.analysis_results:
                    if record.get("analysis_id") == aid:
                        full_result = record.get("full_result", {})
                        full_result["analysis_id"] = aid
                        results.append(full_result)
                        found = True
                        break
                if not found:
                    missing_ids.append(aid)

            if missing_ids:
                return json.dumps({
                    "status": "error",
                    "message": f"Analysis IDs not found: {missing_ids}"
                })

            # Synthesize via the standalone function
            counter = len(self.orch.active_knowledge) + 1
            try:
                entry = _synthesize(
                    results, focus,
                    model=self.orch.model,
                    knowledge_id=f"knowledge_{counter:03d}",
                    synthesis_type=synthesis_type,
                )
            except (ValueError, RuntimeError) as e:
                return json.dumps({"status": "error", "message": str(e)})

            entry["source_analyses"] = analysis_ids
            self.orch.active_knowledge.append(entry)

            # Save to disk
            knowledge_dir = self.orch.base_dir / "knowledge"
            knowledge_dir.mkdir(parents=True, exist_ok=True)
            knowledge_file = knowledge_dir / f"{entry['id']}.json"
            with open(knowledge_file, 'w') as f:
                json.dump(entry, f, indent=2)

            response = {
                "status": "success",
                "knowledge_id": entry["id"],
                "focus": focus,
                "synthesis_type": synthesis_type,
                "summary": entry["summary"],
                "key_findings": entry["key_findings"],
                "saved_to": str(knowledge_file),
                "note": "This knowledge will be automatically injected into all subsequent run_analysis calls."
            }

            # Check if any graduated skill is linked to knowledge with same focus
            for skill_name, source_ids in self.orch._graduated_skill_sources.items():
                for kid in source_ids:
                    for k in self.orch.active_knowledge:
                        if k.get("id") == kid and k.get("focus", "").lower() == focus.lower():
                            response["skill_update_suggested"] = skill_name
                            response["skill_update_note"] = (
                                f"Graduated skill '{skill_name}' is linked to knowledge "
                                f"with the same focus area. Consider calling update_skill "
                                f"to incorporate the new findings."
                            )
                            break
                    if "skill_update_suggested" in response:
                        break
                if "skill_update_suggested" in response:
                    break

            return json.dumps(response)

        self._register_tool(
            func=synthesize_knowledge,
            name="synthesize_knowledge",
            description=(
                "Distill findings from completed analyses into reusable knowledge. "
                "Use when the user wants to learn from reference spectra, derive calibration, "
                "build a reference model, detect trends, learn from failures, or compare methods. "
                "The synthesized knowledge is automatically "
                "injected into all subsequent run_analysis calls as prior knowledge context."
            ),
            parameters={
                "analysis_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of analysis IDs to synthesize knowledge from"
                },
                "focus": {
                    "type": "string",
                    "description": "What to extract/learn (e.g., 'peak assignments for Ti 2p XPS', 'baseline behavior in DSC curves')"
                },
                "synthesis_type": {
                    "type": "string",
                    "enum": ["reference", "trend", "failure", "method"],
                    "description": (
                        "Type of synthesis: 'reference' (calibration/reference extraction, default), "
                        "'trend' (cross-sample trend detection), "
                        "'failure' (failure pattern learning), "
                        "'method' (method selection heuristics)"
                    )
                }
            },
            required=["analysis_ids", "focus"]
        )

        # =====================================================================
        # 14. LIST KNOWLEDGE
        # =====================================================================
        def list_knowledge() -> str:
            """List all active knowledge entries."""
            print(f"  ⚡ Tool: Listing active knowledge...")

            if not self.orch.active_knowledge:
                return json.dumps({
                    "status": "success",
                    "message": "No active knowledge entries.",
                    "entries": []
                })

            entries = []
            for entry in self.orch.active_knowledge:
                entries.append({
                    "id": entry["id"],
                    "focus": entry["focus"],
                    "source_count": len(entry.get("source_analyses", [])),
                    "findings_count": len(entry.get("key_findings", [])),
                    "timestamp": entry.get("timestamp")
                })

            return json.dumps({
                "status": "success",
                "total_entries": len(entries),
                "entries": entries
            })

        self._register_tool(
            func=list_knowledge,
            name="list_knowledge",
            description="Show all active knowledge entries synthesized from previous analyses.",
            parameters={},
            required=[]
        )

        # =====================================================================
        # 15. CLEAR KNOWLEDGE
        # =====================================================================
        def clear_knowledge(knowledge_id: str = None) -> str:
            """Remove active knowledge entries. If knowledge_id is None, removes all."""
            print(f"  ⚡ Tool: Clearing knowledge...")

            knowledge_dir = self.orch.base_dir / "knowledge"

            if knowledge_id is None:
                count = len(self.orch.active_knowledge)
                self.orch.active_knowledge.clear()
                # Remove all files
                if knowledge_dir.exists():
                    for f in knowledge_dir.glob("knowledge_*.json"):
                        f.unlink()
                return json.dumps({
                    "status": "success",
                    "message": f"Cleared all {count} knowledge entries."
                })

            # Find and remove specific entry
            for i, entry in enumerate(self.orch.active_knowledge):
                if entry["id"] == knowledge_id:
                    self.orch.active_knowledge.pop(i)
                    # Remove disk file
                    knowledge_file = knowledge_dir / f"{knowledge_id}.json"
                    if knowledge_file.exists():
                        knowledge_file.unlink()
                    return json.dumps({
                        "status": "success",
                        "message": f"Removed knowledge entry: {knowledge_id}"
                    })

            return json.dumps({
                "status": "error",
                "message": f"Knowledge ID not found: {knowledge_id}"
            })

        self._register_tool(
            func=clear_knowledge,
            name="clear_knowledge",
            description=(
                "Remove active knowledge entries. Specify a knowledge_id to remove a "
                "specific entry, or omit to clear all knowledge."
            ),
            parameters={
                "knowledge_id": {
                    "type": "string",
                    "description": "ID of knowledge entry to remove (omit to clear all)"
                }
            },
            required=[]
        )

        # =====================================================================
        # 16. GRADUATE TO SKILL
        # =====================================================================
        def graduate_to_skill(knowledge_id: str, skill_name: str, domain: str = "curve_fitting") -> str:
            """
            Convert a knowledge entry into a reusable skill (.md file).
            The skill is automatically registered for use in subsequent analyses.
            """
            from scilink.agents.exp_agents.instruct import KNOWLEDGE_TO_SKILL_INSTRUCTIONS

            print(f"  ⚡ Tool: Graduating knowledge '{knowledge_id}' to skill '{skill_name}'...")

            # Find the knowledge entry
            knowledge_entry = None
            for entry in self.orch.active_knowledge:
                if entry.get("id") == knowledge_id:
                    knowledge_entry = entry
                    break

            if knowledge_entry is None:
                return json.dumps({
                    "status": "error",
                    "message": f"Knowledge ID not found: {knowledge_id}"
                })

            # Build knowledge text
            knowledge_text = f"**Focus:** {knowledge_entry.get('focus', '')}\n"
            knowledge_text += f"**Summary:** {knowledge_entry.get('summary', '')}\n"
            knowledge_text += "**Key Findings:**\n"
            for finding in knowledge_entry.get("key_findings", []):
                knowledge_text += f"- {finding}\n"

            # Collect source analysis details
            analysis_details_parts = []
            source_ids = knowledge_entry.get("source_analyses", [])
            for aid in source_ids:
                for record in self.orch.analysis_results:
                    if record.get("analysis_id") == aid:
                        full_result = record.get("full_result", {})
                        parts = [f"### Analysis: {aid}"]

                        da = full_result.get("detailed_analysis", "")
                        if da:
                            parts.append(da[:2000])  # Truncate for prompt size

                        fp = full_result.get("fitting_parameters")
                        if fp:
                            parts.append(f"Fitting parameters: {json.dumps(fp, indent=2, default=str)}")

                        hf = full_result.get("human_feedback", {})
                        if isinstance(hf, dict) and hf.get("user_feedback"):
                            parts.append(f"User feedback: {hf['user_feedback']}")

                        analysis_details_parts.append("\n".join(parts))
                        break

            analysis_details = "\n\n".join(analysis_details_parts) if analysis_details_parts else "No source analysis details available."

            # Call LLM to generate skill content
            prompt = KNOWLEDGE_TO_SKILL_INSTRUCTIONS.format(
                skill_name=skill_name,
                domain=domain,
                knowledge_text=knowledge_text,
                analysis_details=analysis_details,
            )

            try:
                response = self.orch.model.generate_content(
                    contents=[prompt],
                    generation_config=None,
                    safety_settings=None,
                )
                skill_content = response.text if hasattr(response, "text") else str(response)
            except Exception as e:
                return json.dumps({"status": "error", "message": f"LLM call failed: {e}"})

            # Save skill file
            skill_dir = self.orch.base_dir / "graduated_skills"
            skill_dir.mkdir(parents=True, exist_ok=True)
            skill_path = skill_dir / f"{skill_name}.md"
            skill_path.write_text(skill_content)

            # Register the skill
            self.orch.register_skill(str(skill_path))

            # Track the link
            self.orch._graduated_skill_sources[skill_name] = [knowledge_id]

            return json.dumps({
                "status": "success",
                "skill_name": skill_name,
                "skill_path": str(skill_path),
                "source_knowledge_id": knowledge_id,
                "note": f"Skill '{skill_name}' has been registered and will be available in run_analysis."
            })

        self._register_tool(
            func=graduate_to_skill,
            name="graduate_to_skill",
            description=(
                "Convert a knowledge entry into a reusable skill (.md file). "
                "The skill is organized into 5 sections (overview, planning, analysis, "
                "interpretation, validation) and automatically registered for use in "
                "subsequent analyses."
            ),
            parameters={
                "knowledge_id": {
                    "type": "string",
                    "description": "ID of the knowledge entry to graduate"
                },
                "skill_name": {
                    "type": "string",
                    "description": "Name for the new skill (used as filename and reference)"
                },
                "domain": {
                    "type": "string",
                    "description": "Domain/technique area (e.g., 'curve_fitting', 'xps', 'raman'). Default: 'curve_fitting'"
                }
            },
            required=["knowledge_id", "skill_name"]
        )

        # =====================================================================
        # 17. UPDATE SKILL
        # =====================================================================
        def update_skill(skill_name: str, knowledge_ids: list = None) -> str:
            """
            Update a graduated skill with new knowledge entries.
            Preserves the old version as {name}.prev.md.
            """
            from scilink.agents.exp_agents.instruct import SKILL_UPDATE_INSTRUCTIONS

            print(f"  ⚡ Tool: Updating skill '{skill_name}'...")

            # Find the existing skill file
            skill_dir = self.orch.base_dir / "graduated_skills"
            skill_path = skill_dir / f"{skill_name}.md"
            if not skill_path.exists():
                return json.dumps({
                    "status": "error",
                    "message": f"Graduated skill not found: {skill_name}"
                })

            existing_skill = skill_path.read_text()

            # Determine source knowledge IDs
            tracked_ids = self.orch._graduated_skill_sources.get(skill_name, [])
            if knowledge_ids:
                new_ids = knowledge_ids
            else:
                # Use all knowledge entries with matching focus
                focus_areas = set()
                for kid in tracked_ids:
                    for k in self.orch.active_knowledge:
                        if k.get("id") == kid:
                            focus_areas.add(k.get("focus", "").lower())
                new_ids = [
                    k["id"] for k in self.orch.active_knowledge
                    if k["id"] not in tracked_ids and k.get("focus", "").lower() in focus_areas
                ]

            if not new_ids:
                return json.dumps({
                    "status": "error",
                    "message": "No new knowledge entries found to update the skill with."
                })

            # Collect new knowledge texts
            new_knowledge_parts = []
            for kid in new_ids:
                for k in self.orch.active_knowledge:
                    if k.get("id") == kid:
                        part = f"### {kid}\n**Focus:** {k.get('focus', '')}\n"
                        part += f"**Summary:** {k.get('summary', '')}\n"
                        part += "**Key Findings:**\n"
                        for f in k.get("key_findings", []):
                            part += f"- {f}\n"
                        new_knowledge_parts.append(part)
                        break

            new_knowledge = "\n\n".join(new_knowledge_parts)

            # Call LLM to produce updated skill
            prompt = SKILL_UPDATE_INSTRUCTIONS.format(
                skill_name=skill_name,
                existing_skill=existing_skill,
                new_knowledge=new_knowledge,
            )

            try:
                response = self.orch.model.generate_content(
                    contents=[prompt],
                    generation_config=None,
                    safety_settings=None,
                )
                updated_content = response.text if hasattr(response, "text") else str(response)
            except Exception as e:
                return json.dumps({"status": "error", "message": f"LLM call failed: {e}"})

            # Save previous version
            prev_path = skill_dir / f"{skill_name}.prev.md"
            prev_path.write_text(existing_skill)

            # Write updated skill
            skill_path.write_text(updated_content)

            # Update source tracking
            all_ids = list(set(tracked_ids + new_ids))
            self.orch._graduated_skill_sources[skill_name] = all_ids

            # Re-register the skill
            self.orch.register_skill(str(skill_path))

            return json.dumps({
                "status": "success",
                "skill_name": skill_name,
                "skill_path": str(skill_path),
                "previous_version": str(prev_path),
                "new_knowledge_ids": new_ids,
                "total_source_ids": all_ids,
                "note": f"Skill '{skill_name}' has been updated. Previous version saved as {prev_path.name}."
            })

        self._register_tool(
            func=update_skill,
            name="update_skill",
            description=(
                "Update a graduated skill with new knowledge entries. "
                "Use when new knowledge has been synthesized and a linked skill "
                "should incorporate the new findings. The old version is preserved."
            ),
            parameters={
                "skill_name": {
                    "type": "string",
                    "description": "Name of the graduated skill to update"
                },
                "knowledge_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Specific knowledge IDs to incorporate (omit to auto-detect from matching focus area)"
                }
            },
            required=["skill_name"]
        )

        # =====================================================================
        # 18. SAVE FILE
        # =====================================================================
        def save_file(filename: str, content: str, subfolder: str = "") -> str:
            """
            Save text content (reports, summaries, tables, scripts) to a file
            in the session directory.
            """
            print(f"  ⚡ Tool: Saving file '{filename}'...")

            # Sanitise: strip path separators from filename to prevent traversal.
            safe_name = Path(filename).name
            if not safe_name:
                return json.dumps({
                    "status": "error",
                    "message": "Invalid filename.",
                })

            target_dir = self.orch.base_dir
            if subfolder:
                safe_sub = Path(subfolder).name
                target_dir = target_dir / safe_sub
            target_dir.mkdir(parents=True, exist_ok=True)
            dest = target_dir / safe_name

            try:
                dest.write_text(content, encoding="utf-8")
                print(f"    💾 Saved: {dest}")
                return json.dumps({
                    "status": "success",
                    "path": str(dest),
                    "size_bytes": dest.stat().st_size,
                })
            except Exception as e:
                logging.error(f"save_file failed: {e}")
                return json.dumps({
                    "status": "error",
                    "message": str(e),
                })

        self._register_tool(
            func=save_file,
            name="save_file",
            description=(
                "Save text content (reports, summaries, tables, scripts, notes) "
                "to a file in the session directory. Use this to persist "
                "synthesized knowledge summaries, analysis reports, exported "
                "results, or any text artifact the user requests."
            ),
            parameters={
                "filename": {
                    "type": "string",
                    "description": (
                        "Name of the file to create, e.g. 'analysis_report.md', "
                        "'peak_positions.csv', or 'summary.txt'."
                    ),
                },
                "content": {
                    "type": "string",
                    "description": "The text content to write to the file.",
                },
                "subfolder": {
                    "type": "string",
                    "description": (
                        "Optional subfolder within the session directory, "
                        "e.g. 'reports' or 'exports'. Created if it doesn't exist."
                    ),
                },
            },
            required=["filename", "content"]
        )

        # =====================================================================
        # READ DOCUMENT
        # =====================================================================
        def read_document(paths) -> str:
            """Read one or more PDF/DOCX/MD/TXT documents; return the combined
            text and persist a literature_file for run_analysis."""
            if isinstance(paths, str):
                paths = [paths]
            if not paths:
                return json.dumps({
                    "status": "error",
                    "message": "No document path provided.",
                })
            print(f"  📄 Tool: Reading {len(paths)} document(s)...")
            docs, errors = [], []
            for p in paths:
                dp = Path(p)
                if not dp.is_file():
                    errors.append(f"Not a file: {p}")
                    continue
                try:
                    docs.append((dp, _extract_document_text(dp)))
                except ValueError as e:
                    errors.append(str(e))
                except Exception as e:
                    logging.error(f"read_document failed for {p}: {e}")
                    errors.append(f"Could not read {dp.name}: {e}")
            if not docs:
                return json.dumps({
                    "status": "error",
                    "message": "No documents could be read.",
                    "errors": errors,
                })
            combined = "\n\n---\n\n".join(
                f"## {dp.name}\n\n{info['text']}" for dp, info in docs
            )
            combined_truncated = len(combined) > _READ_DOC_MAX_CHARS
            if combined_truncated:
                combined = combined[:_READ_DOC_MAX_CHARS]
            # Persist as a literature_file so run_analysis can ground its plan
            # in these documents — the same channel search_literature uses.
            lit_path = None
            try:
                lit_dir = self.orch.base_dir / "literature"
                lit_dir.mkdir(parents=True, exist_ok=True)
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                lit_path = lit_dir / f"provided_documents_{ts}.md"
                lit_path.write_text(combined)
            except Exception as e:
                logging.error(f"read_document: could not save literature file: {e}")
            return json.dumps({
                "status": "success",
                "file_path": str(lit_path) if lit_path else None,
                "n_documents": len(docs),
                "documents": [
                    {"name": dp.name,
                     **{k: v for k, v in info.items() if k != "text"}}
                    for dp, info in docs
                ],
                "errors": errors or None,
                "combined_truncated": combined_truncated,
                "text": combined,
                "hint": (
                    "Pass file_path as `literature_file` to the next "
                    "run_analysis() call so the planner produces a "
                    "document-informed plan."
                ),
            })

        self._register_tool(
            func=read_document,
            name="read_document",
            description=(
                "Read one or more documents the user provided — PDF, DOCX, "
                "Markdown, or text files (a methods paper, protocol, prior "
                "report, notes). Returns the combined text AND saves a "
                "literature file: pass the returned `file_path` as "
                "`literature_file` to run_analysis() so the provided documents "
                "drive the analysis plan, exactly as a search_literature "
                "result does. For a handful of documents read straight into "
                "context — it runs NO external literature search and builds no "
                "index (use search_literature for the wider literature). Pass "
                "absolute file paths."
            ),
            parameters={
                "paths": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": (
                        "Absolute path(s) to the document(s) to read (.pdf, "
                        ".docx, .md, or .txt). Multiple documents are combined "
                        "into one literature file."
                    ),
                },
            },
            required=["paths"]
        )

    def _register_tool(
        self,
        func: Callable,
        name: str,
        description: str,
        parameters: Dict[str, Any],
        required: list = None
    ):
        """Register a tool in OpenAI format."""
        self.functions_map[name] = func
        
        openai_schema = {
            "type": "function",
            "function": {
                "name": name,
                "description": description,
                "parameters": {
                    "type": "object",
                    "properties": parameters,
                    "required": required or []
                }
            }
        }
        self.openai_schemas.append(openai_schema)

    def _update_skill_description(self, custom_skills: dict) -> None:
        """Update the ``skill`` parameter description in ``run_analysis``
        to include newly registered custom skills."""
        for schema in self.openai_schemas:
            fn = schema.get("function", {})
            if fn.get("name") != "run_analysis":
                continue
            skill_prop = fn["parameters"]["properties"].get("skill")
            if skill_prop is None:
                break
            skill_prop["description"] = _build_skill_description(
                getattr(self.orch, "_agent_registry", None),
                custom_skills,
            )
            break

    def execute_tool(self, tool_name: str, **kwargs) -> str:
        """Execute a tool by name with given arguments."""
        if tool_name not in self.functions_map:
            return json.dumps({
                "status": "error",
                "message": f"Tool '{tool_name}' not found"
            })
        
        try:
            return self.functions_map[tool_name](**kwargs)
        except Exception as e:
            logging.error(f"Tool execution error ({tool_name}): {e}", exc_info=True)
            return json.dumps({
                "status": "error",
                "message": str(e),
                "tool": tool_name
            })