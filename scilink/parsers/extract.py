"""Unified lightweight document text extraction (Tier 1).

``extract_text`` turns a single document into flat text — for callers that
want a handful of documents read straight into context, with no chunking,
embeddings, or retrieval. The heavyweight ingestion/retrieval path lives in
``scilink.knowledge``.
"""

from pathlib import Path
from typing import Any, Dict, Union

from .pdf_parser import _extract_pdf_blocks, _assemble_flat_text
from .ocr import DEFAULT_OCR_DPI, MAX_OCR_PAGES


def extract_text(path: Union[str, Path], max_pages: int = None,
                 ocr_model: Any = None, ocr_dpi: int = DEFAULT_OCR_DPI,
                 max_ocr_pages: int = MAX_OCR_PAGES) -> Dict[str, Any]:
    """Extract plain text (and markdown tables, for PDFs) from a document.

    Supports text-like (``.pdf``, ``.docx``, ``.md``, ``.txt``, ``.json``,
    ``.yaml``/``.yml``) and tabular (``.csv``, ``.xlsx``/``.xls``) documents.
    Returns a dict with ``text``, ``n_chars`` and a format-specific count
    (``n_pages`` for PDFs, ``n_paragraphs`` for DOCX). PDFs also report
    ``n_ocr_pages`` — pages transcribed via the vision-OCR fallback.
    Tabular files are previewed via the adaptive Excel parser — small files
    yield the full table as Markdown, large files a statistical summary; an
    auto-detected sibling JSON metadata file (e.g. ``data.json`` next to
    ``data.xlsx``) enriches the preview with title / objective / column
    definitions.

    No truncation is applied — any length cap is the caller's policy.

    Args:
        path: Path to the document.
        max_pages: For PDFs only — stop after this many pages and skip table
            extraction. A lightweight mode for previews/probes; ``n_pages``
            still reports the true total.
        ocr_model: Optional vision LLM (any object with ``generate_content``).
            When given, scanned/sparse PDF pages are transcribed via OCR;
            without it, such pages simply yield their (sparse) text.
        ocr_dpi: Render resolution for OCR'd pages.
        max_ocr_pages: Cap on the number of pages sent to vision-OCR.

    Raises:
        ValueError: For an unsupported extension.
        ImportError: For a ``.docx`` when ``python-docx`` is not installed.
    """
    path = Path(path)
    ext = path.suffix.lower()
    info: Dict[str, Any] = {}

    if ext == ".pdf":
        page_texts, table_chunks, n_pages, ocr_pages = _extract_pdf_blocks(
            str(path), max_pages=max_pages, ocr_model=ocr_model,
            ocr_dpi=ocr_dpi, max_ocr_pages=max_ocr_pages,
        )
        text = _assemble_flat_text(page_texts, table_chunks, ocr_pages)
        info["n_pages"] = n_pages
        info["n_ocr_pages"] = len(ocr_pages)
    elif ext == ".docx":
        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "Reading .docx documents requires python-docx "
                "(`pip install python-docx`)."
            ) from e
        d = docx.Document(str(path))
        info["n_paragraphs"] = len(d.paragraphs)
        text = "\n".join(p.text for p in d.paragraphs)
    elif ext in (".md", ".txt", ".json", ".yaml", ".yml"):
        text = path.read_text(errors="replace")
    elif ext in (".csv", ".xlsx", ".xls"):
        from .excel_parser import parse_adaptive_excel
        # Auto-discover a sibling JSON metadata file (e.g. data.xlsx +
        # data.json with title / objective / column_definitions) — the same
        # convention parse_adaptive_excel uses elsewhere in the codebase.
        sidecar = path.with_suffix(".json")
        chunks = parse_adaptive_excel(
            str(path),
            context_path=str(sidecar) if sidecar.exists() else None,
        )
        if chunks:
            summary = next(
                (c for c in chunks if c["metadata"].get("content_type")
                 in ("dataset_summary", "dataset_package")),
                chunks[0],
            )
            text = summary["text"]
        else:
            text = ""
    else:
        raise ValueError(
            f"Unsupported document type '{ext}' — extract_text handles "
            f".pdf, .docx, .md, .txt, .json, .yaml/.yml, "
            f".csv and .xlsx/.xls."
        )

    text = text.strip()
    info["text"] = text
    info["n_chars"] = len(text)
    return info
